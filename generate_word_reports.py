#!/usr/bin/env python3
"""Generate comprehensive Word reports for three antibodies."""

import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from scripts.humanize.evaluate_humanness import (
    evaluate_sequence, evaluate_sequence_imgt
)
from scripts.humanize.imgt_numbering import (
    number_with_abrsa_imgt, load_imgt_germline, compare_imgt_posmaps_direct,
    IMGT_REGIONS
)

# Three antibodies to evaluate
antibodies = {
    "C45H4_VH": "QVQLVQSGAEVKKPGASVKVSCKASGYTFTSYWMHWVRQAPGQGLEWIGQIDPSDSYTYYNEDFKDRATLTVDQSTSTAYMELSSLRSEDTAVYYCAKGYYDYDWGYAMDYWGQGTTVTVSS",
    "E51H2_VH": "QVQLVQSGAEVKKPGSSVKVSCKDSDGTVFPIAYMSWVRQAPGQGLEWMGGIFPSIGRTIYGEKFEDRVTITADESTSTAYMELSSLRSEDTAVYYCARGRTYWEYYHAMDNWGQGTTVTVSS",
    "E51H6_VH": "QVQLVQSGAEVKKPGSSVKVSCKDSDSEVFPIAYMSWVRQAPGQGFEWIGGIFPSIGRTIYGEKFEDRATLDADTSTNTAYMELSSLRSEDTAVYYCARGRTYWEYYHAMDNWGQGTTVTVSS",
}

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color
    })
    shading_elm.append(shading)

def get_usan_class(fr_identity):
    """Get academic humanization estimate based on FR identity.
    
    NOTE: This is an ACADEMIC ESTIMATE based on sequence identity.
    The official USAN/INN naming is based on the technology used to create
    the antibody (transgenic mouse, CDR grafting, phage display, etc.),
    NOT on sequence identity thresholds.
    """
    if fr_identity >= 0.95:
        return "Academic: likely Human-like"
    elif fr_identity >= 0.85:
        return "Academic: likely Humanized-like"
    elif fr_identity >= 0.70:
        return "Academic: likely Chimeric-like"
    else:
        return "Academic: likely Murine-like"

def get_imgt_region(pos_label):
    """Get IMGT region from position label like 'H105' or 'H111A'."""
    from scripts.humanize.imgt_numbering import IMGT_REGIONS
    
    # Extract numeric position
    num_str = ""
    for c in pos_label[1:]:
        if c.isdigit():
            num_str += c
        else:
            break
    
    if not num_str:
        return "Unknown"
    
    num = int(num_str)
    for region, (start, end) in IMGT_REGIONS.items():
        if start <= num <= end:
            return region
    return "Unknown"

def get_cdr_sequences(sequence):
    """Get CDR sequences using AbRSA IMGT numbering."""
    numbered = number_with_abrsa_imgt(sequence, 'H')
    if not numbered:
        return None
    
    posmap = numbered.posmap()
    
    cdrs = {}
    for region, (start, end) in IMGT_REGIONS.items():
        if region.startswith('CDR'):
            seq = ''
            for num in range(start, end + 1):
                # Add base position
                label = f'H{num}'
                aa = posmap.get(label, '-')
                if aa != '-':
                    seq += aa
                
                # Add insertions (e.g., H111A, H111B, etc.)
                for letter in 'ABCDEFGHIJK':
                    ins_label = f'H{num}{letter}'
                    ins_aa = posmap.get(ins_label, None)
                    if ins_aa:
                        seq += ins_aa
            cdrs[region] = seq
    
    return cdrs

def get_kabat_cdr_sequences(sequence):
    """Get CDR sequences using Kabat numbering."""
    from scripts.humanize.numbering import number_heavy
    
    numbered = number_heavy(sequence)
    posmap = numbered.posmap()
    
    # Kabat CDR regions for VH
    kabat_cdrs = {
        'CDR1': ('H31', 'H35'),
        'CDR2': ('H50', 'H65'),
        'CDR3': ('H95', 'H102'),
    }
    
    def extract_cdr_seq(posmap, start_label, end_label):
        """Extract CDR sequence including insertions."""
        # Extract numeric positions
        start_num = int(''.join(c for c in start_label if c.isdigit()))
        end_num = int(''.join(c for c in end_label if c.isdigit()))
        
        seq = ''
        for num in range(start_num, end_num + 1):
            # Add base position
            label = f'H{num}'
            aa = posmap.get(label, '-')
            if aa != '-':
                seq += aa
            
            # Add insertions (e.g., H100A, H100B, etc.)
            for letter in 'ABCDEFGHIJK':
                ins_label = f'H{num}{letter}'
                ins_aa = posmap.get(ins_label, None)
                if ins_aa:
                    seq += ins_aa
        
        return seq
    
    cdrs = {}
    for region, (start, end) in kabat_cdrs.items():
        cdrs[region] = extract_cdr_seq(posmap, start, end)
    
    return cdrs

def evaluate_antibody(name, sequence):
    """Evaluate a single antibody and return comprehensive results."""
    results = {}
    
    # Kabat evaluation
    scored_kabat, numbered_kabat = evaluate_sequence("H", sequence)
    best_gene_kabat, best_scores_kabat = scored_kabat[0]
    
    # Count differences
    q_map = numbered_kabat.posmap()
    g_map = best_gene_kabat.numbered.posmap() if best_gene_kabat.numbered else {}
    diffs = []
    for pos in q_map:
        if pos in g_map and q_map[pos] != g_map[pos]:
            region = numbered_kabat.region_of(pos) or "?"
            diffs.append((pos, q_map[pos], g_map[pos], region))
    
    results['kabat'] = {
        'best_gene': best_gene_kabat.gene_id,
        'fr_identity': best_scores_kabat['fr_identity'],
        'cdr_identity': best_scores_kabat['cdr_identity'],
        'all_identity': best_scores_kabat['all_identity'],
        'n_diff': len(diffs),
        'regions': {
            'FR1': best_scores_kabat.get('fr1_identity', 0),
            'FR2': best_scores_kabat.get('fr2_identity', 0),
            'FR3': best_scores_kabat.get('fr3_identity', 0),
            'CDR1': best_scores_kabat.get('cdr1_identity', 0),
            'CDR2': best_scores_kabat.get('cdr2_identity', 0),
        },
        'diffs': diffs,
    }
    
    # IMGT evaluation
    scored_imgt, numbered_imgt = evaluate_sequence_imgt("H", sequence)
    
    if scored_imgt:
        best_gene_imgt, best_scores_imgt = scored_imgt[0]
        
        # Get region stats
        regions_imgt = best_scores_imgt.get('imgt_region_stats', {})
        
        # Count IMGT differences
        imgt_diffs = []
        if hasattr(numbered_imgt, 'posmap'):
            q_map_imgt = numbered_imgt.posmap()
            g_map_imgt = best_gene_imgt._imgt_posmap if hasattr(best_gene_imgt, '_imgt_posmap') else {}
            for pos in q_map_imgt:
                if pos in g_map_imgt and q_map_imgt[pos] != g_map_imgt[pos]:
                    region = get_imgt_region(pos)
                    imgt_diffs.append((pos, q_map_imgt[pos], g_map_imgt[pos], region))
        
        results['imgt'] = {
            'best_gene': best_gene_imgt.gene_id,
            'fr_identity': best_scores_imgt['fr_identity'],
            'cdr_identity': best_scores_imgt['cdr_identity'],
            'all_identity': best_scores_imgt['all_identity'],
            'n_diff': len(imgt_diffs),
            'usan_class': get_usan_class(best_scores_imgt['fr_identity']),
            'regions': regions_imgt,
            'diffs': imgt_diffs,
        }
    else:
        results['imgt'] = {
            'best_gene': 'N/A',
            'fr_identity': 0,
            'cdr_identity': 0,
            'all_identity': 0,
            'n_diff': 0,
            'usan_class': 'N/A',
            'regions': {},
            'diffs': [],
        }
    
    # Get CDR sequences (both IMGT and Kabat)
    results['cdrs'] = get_cdr_sequences(sequence)
    results['kabat_cdrs'] = get_kabat_cdr_sequences(sequence)
    
    return results

def create_summary_table(doc, antibody_results):
    """Create summary table at the beginning of the document."""
    doc.add_heading('Summary Table', level=1)
    
    # Kabat summary table
    doc.add_heading('Kabat Numbering', level=2)
    table_kabat = doc.add_table(rows=1, cols=7)
    table_kabat.style = 'Table Grid'
    table_kabat.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    headers = ['Antibody', 'Best Germline', 'FR id', 'CDR id', 'Full id', '# Diff', 'Estimate']
    for i, header in enumerate(headers):
        cell = table_kabat.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    for name, results in antibody_results.items():
        kabat = results['kabat']
        row = table_kabat.add_row()
        row.cells[0].text = name
        row.cells[1].text = kabat['best_gene']
        row.cells[2].text = f"{kabat['fr_identity']*100:.1f}%"
        row.cells[3].text = f"{kabat['cdr_identity']*100:.1f}%"
        row.cells[4].text = f"{kabat['all_identity']*100:.1f}%"
        row.cells[5].text = str(kabat['n_diff'])
        row.cells[6].text = "N/A"
    
    doc.add_paragraph()
    
    # IMGT summary table
    doc.add_heading('IMGT Numbering (WHO/INN/USAN Standard)', level=2)
    table_imgt = doc.add_table(rows=1, cols=7)
    table_imgt.style = 'Table Grid'
    table_imgt.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    for i, header in enumerate(headers):
        cell = table_imgt.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '548235')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    for name, results in antibody_results.items():
        imgt = results['imgt']
        row = table_imgt.add_row()
        row.cells[0].text = name
        row.cells[1].text = imgt['best_gene']
        row.cells[2].text = f"{imgt['fr_identity']*100:.1f}%"
        row.cells[3].text = f"{imgt['cdr_identity']*100:.1f}%"
        row.cells[4].text = f"{imgt['all_identity']*100:.1f}%"
        row.cells[5].text = str(imgt['n_diff'])
        row.cells[6].text = imgt['usan_class']
    
    doc.add_paragraph()

def add_antibody_section(doc, name, sequence, results):
    """Add detailed section for one antibody."""
    doc.add_heading(f'{name}', level=1)
    
    # Sequence info
    doc.add_paragraph(f'Sequence: {sequence}')
    doc.add_paragraph(f'Length: {len(sequence)} amino acids')
    
    # CDR sequences (Kabat)
    if results.get('kabat_cdrs'):
        doc.add_heading('CDR Sequences (Kabat)', level=2)
        for region, seq in results['kabat_cdrs'].items():
            doc.add_paragraph(f'{region}: {seq}', style='List Bullet')
    
    # CDR sequences (IMGT)
    if results['cdrs']:
        doc.add_heading('CDR Sequences (IMGT)', level=2)
        for region, seq in results['cdrs'].items():
            doc.add_paragraph(f'{region}: {seq}', style='List Bullet')
    
    # Kabat results
    doc.add_heading('Kabat Numbering', level=2)
    kabat = results['kabat']
    
    doc.add_paragraph(f'Best match: {kabat["best_gene"]}')
    doc.add_paragraph(f'FR identity: {kabat["fr_identity"]*100:.1f}%')
    doc.add_paragraph(f'CDR identity: {kabat["cdr_identity"]*100:.1f}%')
    doc.add_paragraph(f'Overall identity: {kabat["all_identity"]*100:.1f}%')
    
    # Region breakdown
    doc.add_heading('Region Breakdown', level=3)
    for region, identity in kabat['regions'].items():
        doc.add_paragraph(f'{region}: {identity*100:.1f}%', style='List Bullet')
    
    # Differing positions
    if kabat['diffs']:
        doc.add_heading('Differing Positions', level=3)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        headers = ['Position', 'Query', 'Germline', 'Region']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        
        for pos, q_aa, g_aa, region in kabat['diffs'][:30]:
            row = table.add_row()
            row.cells[0].text = pos
            row.cells[1].text = q_aa
            row.cells[2].text = g_aa
            row.cells[3].text = region
    
    # IMGT results
    doc.add_heading('IMGT Numbering', level=2)
    imgt = results['imgt']
    
    doc.add_paragraph(f'Best match: {imgt["best_gene"]}')
    doc.add_paragraph(f'FR identity: {imgt["fr_identity"]*100:.1f}%')
    doc.add_paragraph(f'CDR identity: {imgt["cdr_identity"]*100:.1f}%')
    doc.add_paragraph(f'Overall identity: {imgt["all_identity"]*100:.1f}%')
    doc.add_paragraph(f'Humanization estimate: {imgt["usan_class"]}')
    doc.add_paragraph(f'(Note: Official USAN/INN naming requires knowledge of manufacturing process)', style='Intense Quote')
    
    # Region breakdown
    doc.add_heading('IMGT Region Breakdown', level=3)
    for region, stats in imgt['regions'].items():
        doc.add_paragraph(f'{region}: {stats["identity"]*100:.1f}% ({stats["match"]}/{stats["count"]})', style='List Bullet')
    
    # Differing positions
    if imgt['diffs']:
        doc.add_heading('IMGT Differing Positions', level=3)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        headers = ['Position', 'Query', 'Germline', 'Region']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        
        for pos, q_aa, g_aa, region in imgt['diffs'][:30]:
            row = table.add_row()
            row.cells[0].text = pos
            row.cells[1].text = q_aa
            row.cells[2].text = g_aa
            row.cells[3].text = region
    
    doc.add_page_break()

def main():
    # Create output directory
    os.makedirs("outputs/reports", exist_ok=True)
    
    print("正在评估抗体...")
    
    # Evaluate and generate individual report for each antibody
    for name, sequence in antibodies.items():
        print(f"  评估 {name}...")
        results = evaluate_antibody(name, sequence)
        
        # Create individual Word document
        print(f"  生成 {name} 报告...")
        doc = Document()
        
        # Title
        title = doc.add_heading(f'{name} - Humanization Evaluation Report', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Overview
        doc.add_heading('Overview', level=1)
        doc.add_paragraph(f'Antibody: {name}')
        doc.add_paragraph(f'Sequence: {sequence}')
        doc.add_paragraph(f'Length: {len(sequence)} amino acids')
        
        # Evaluation Criteria
        doc.add_heading('Evaluation Criteria', level=2)
        doc.add_paragraph('Kabat numbering: FR1 (1-30) + CDR1 (31-35) + FR2 (36-49) + CDR2 (50-65) + FR3 (66-94) + CDR3 (95-102) + FR4 (103-113)')
        doc.add_paragraph('IMGT numbering: FR1 (1-26) + CDR1 (27-38) + FR2 (39-55) + CDR2 (56-65) + FR3 (66-104) + CDR3 (105-117) + FR4 (118-128)')
        doc.add_paragraph('Excluded: CDR3 (from donor antibody) and FR4 (from human J gene)')
        doc.add_paragraph('Humanization estimate (Academic convention): Based on sequence identity, not official USAN/INN naming')
        doc.add_paragraph('Note: Official USAN/INN naming is based on the technology used (transgenic mouse, CDR grafting, phage display, etc.)')
        
        # Add antibody section
        add_antibody_section(doc, name, sequence, results)
        
        # Save individual document
        output_path = f"outputs/reports/{name}_report.docx"
        doc.save(output_path)
        print(f"  报告已保存: {output_path}")
    
    print("\n所有报告生成完成！")

if __name__ == "__main__":
    main()
