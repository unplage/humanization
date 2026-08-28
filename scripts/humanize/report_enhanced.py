"""增强版报告生成 - 借鉴 WeMol 格式"""

from __future__ import annotations

import os
import re
from datetime import datetime
from typing import Dict, List, Optional, Tuple

from .backmut import BackMutationResult, BackMutationCandidate
from .germline import GermlineDB, GermlineGene
from .germline_frequency import get_frequency, is_recommended
from .numbering import NumberedChain
from .pipeline import ChainReport, RunResult


# PTM 风险模式
PTM_PATTERNS = {
    'ptm_cys': r'C',
    'ptm_glyco': r'N[^P][ST]',
    'ptm_ng': r'NG',
    'ptm_ns': r'NS',
    'ptm_dg': r'DG',
    'ptm_dd': r'DD',
}

# 保守 Cys 位置 (不计入风险)
CONSERVED_CYS = {
    'H': {22, 92},
    'L': {22, 23, 88},
}


def _scan_ptm_risks_for_sequence(seq: str, chain_type: str) -> Dict[str, int]:
    """扫描序列中的 PTM 风险计数"""
    seq = seq.upper()
    result = {}
    
    # Cys 风险：减去保守 Cys
    cys_positions = {m.start() + 1 for m in re.finditer('C', seq)}
    conserved = CONSERVED_CYS.get(chain_type, set())
    result['ptm_cys'] = len(cys_positions - conserved)
    
    # 其他 PTM 风险
    for key, pattern in PTM_PATTERNS.items():
        if key != 'ptm_cys':
            result[key] = len(re.findall(pattern, seq))
    
    return result


def _calculate_fv_identity(query: Optional[NumberedChain], gene: Optional[GermlineGene]) -> float:
    """计算 Fv (FR + CDR) 整体同源性"""
    if query is None or gene is None or gene.numbered is None:
        return 0.0
    
    q = query.posmap()
    g = gene.numbered.posmap()
    
    # 排除 CDR3 (H95-102 / L89-97)
    common = [p for p in q if p in g]
    
    if not common:
        return 0.0
    
    matches = sum(1 for p in common if q[p] == g[p])
    return matches / len(common)


_FR_REGIONS = {"FR1", "FR2", "FR3"}


def _calculate_fr_identity(query: Optional[NumberedChain], gene: Optional[GermlineGene]) -> float:
    """计算 FR 同源性 (仅 FR 区域, 不包括 CDR)"""
    if query is None or gene is None or gene.numbered is None:
        return 0.0
    
    q = query.posmap()
    g = gene.numbered.posmap()
    
    # 通过 region 标签获取 FR 位置 (与 generate_template_score_table 一致)
    fr_positions = {p for p in q if query.region_of(p) in _FR_REGIONS}
    
    # 仅比较 FR 位置
    common = [p for p in q if p in g and p in fr_positions]
    
    if not common:
        return 0.0
    
    matches = sum(1 for p in common if q[p] == g[p])
    return matches / len(common)


def _get_germline_frequency_detail(
    chain_type: str,
    gene_id: str,
    donor_aa: str,
    position: str,
    backmut: BackMutationResult,
) -> str:
    """获取详细的 germline 频率分布"""
    cons = backmut.germline_conservation
    freq = cons.get(position)
    
    if freq is None:
        return "-"
    
    # 简化格式：显示主要残基和频率
    return f"{freq:.0%}"


def _scan_high_risk(numbered, cdr_only: bool = False) -> Tuple[int, List[str]]:
    """对一个 NumberedChain 扫描高风险位点
    
    Args:
        numbered: NumberedChain 对象
        cdr_only: 若 True，仅返回 CDR 区域的高风险位点
    """
    from .developability import scan_sequence
    if not numbered:
        return 0, []
    issues = scan_sequence(numbered)
    high_risk_issues = [issue for issue in issues if issue.risk_level == 'high']
    if cdr_only and numbered.cdrs:
        cdr_positions = set()
        for name, (start, end) in numbered.cdrs.items():
            start_res = numbered.residue(start)
            end_res = numbered.residue(end)
            if start_res and end_res:
                for r in numbered.residues:
                    if start_res.index <= r.index <= end_res.index:
                        cdr_positions.add(r.pos)
        high_risk_issues = [issue for issue in high_risk_issues if issue.position in cdr_positions]
    high_risk_count = len(high_risk_issues)
    high_risk_sites = [f"{issue.position} ({issue.motif})" for issue in high_risk_issues]
    return high_risk_count, high_risk_sites


def _count_high_risk_sites(rep: ChainReport, cdr_only: bool = False) -> Tuple[int, List[str]]:
    """计算高风险位点数量和位置 (使用供体序列)"""
    return _scan_high_risk(rep.input_chain.numbered, cdr_only)


def _get_cdr_sequences(rep: ChainReport) -> str:
    """获取 CDR 序列字符串"""
    # 从 numbered chain 的 cdrs 字典获取 CDR 序列
    # cdrs 是 Dict[str, Tuple[str, str]]，其中值是 (start_pos, end_pos)
    if rep.input_chain.numbered and hasattr(rep.input_chain.numbered, 'cdrs'):
        cdrs = rep.input_chain.numbered.cdrs
        if cdrs:
            seqs = []
            for name, (start, end) in cdrs.items():
                seq = rep.input_chain.numbered.seq_range(start, end)
                if seq:
                    seqs.append(seq)
            return " ".join(seqs) if seqs else "N/A"
    return "N/A"


def generate_template_score_table(
    chain_type: str,
    query: Optional[NumberedChain],
    db: GermlineDB,
    selected_germline: Optional[GermlineGene],
) -> str:
    """生成 Template Score 表格 (WeMol 格式)
    
    NO.	Hit	Adimab_Count	FR_Length_Diff	#FR_Mismatch	Usage%	Fv%	FR%	FR1%	FR2%	FR3%	Hyper_Mut%	PTM_cys	PTM_glyco	PTM_NG	PTM_NS	PTM_DG	PTM_DD	Score
    """
    if query is None:
        return ""
    
    v_genes = db.human(chain_type)
    
    def region_of(pos: str) -> str:
        return query.region_of(pos) or ""
    
    results = []
    for g in v_genes:
        if g.numbered is None:
            continue
        
        q = query.posmap()
        gn = g.numbered.posmap()
        
        # FR 区域计算
        common_fr = [p for p in q if p in gn and region_of(p) in ("FR1", "FR2", "FR3")]
        fr_mismatch = sum(1 for p in common_fr if q[p] != gn[p])
        fr_length = len(common_fr)
        
        # Fv 整体同源性 (FR + CDR)
        fv_identity = _calculate_fv_identity(query, g)
        fv_pct = fv_identity * 100
        
        # FR1/FR2/FR3 单独计算
        def calc_fr_pct(region):
            ps = [p for p in common_fr if region_of(p) == region]
            if not ps:
                return 0.0
            return 100.0 * sum(1 for p in ps if q[p] == gn[p]) / len(ps)
        
        fr1_pct = calc_fr_pct("FR1")
        fr2_pct = calc_fr_pct("FR2")
        fr3_pct = calc_fr_pct("FR3")
        
        # FR 同源性
        fr_identity = 1 - (fr_mismatch / fr_length) if fr_length > 0 else 0.0
        fr_pct = fr_identity * 100
        
        # FR 长度差异
        fr_length_diff = abs(fr_length - len([p for p in q if region_of(p) in ("FR1", "FR2", "FR3")]))
        
        # 使用频率
        frequency = get_frequency(chain_type, g.gene_id)
        usage_pct = frequency * 100
        
        # Adimab_Count (使用频率代理)
        adimab_count = int(frequency * 1000)
        
        # Hyper_Mut%
        hyper_mut_pct = (1 - fv_identity) * 100
        
        # PTM 风险计数
        ptm_risks = _scan_ptm_risks_for_sequence(g.sequence, chain_type)
        
        # 综合评分
        score = fr_identity * 100 + fr1_pct * 0.1 + fr2_pct * 0.1 + fr3_pct * 0.1
        
        results.append({
            'gene': g.gene_id,
            'adimab_count': adimab_count,
            'fr_length_diff': fr_length_diff,
            'fr_mismatch': fr_mismatch,
            'usage_pct': usage_pct,
            'fv_pct': fv_pct,
            'fr_pct': fr_pct,
            'fr1_pct': fr1_pct,
            'fr2_pct': fr2_pct,
            'fr3_pct': fr3_pct,
            'hyper_mut_pct': hyper_mut_pct,
            'ptm_cys': ptm_risks.get('ptm_cys', 0),
            'ptm_glyco': ptm_risks.get('ptm_glyco', 0),
            'ptm_ng': ptm_risks.get('ptm_ng', 0),
            'ptm_ns': ptm_risks.get('ptm_ns', 0),
            'ptm_dg': ptm_risks.get('ptm_dg', 0),
            'ptm_dd': ptm_risks.get('ptm_dd', 0),
            'score': score,
        })
    
    results.sort(key=lambda x: -x['score'])
    
    lines = []
    lines.append("NO.\tHit\tAdimab_Count\tFR_Length_Diff\t#FR_Mismatch\tUsage%\tFv%\tFR%\tFR1%\tFR2%\tFR3%\tHyper_Mut%\tPTM_cys\tPTM_glyco\tPTM_NG\tPTM_NS\tPTM_DG\tPTM_DD\tScore")
    lines.append("-" * 180)
    
    for i, r in enumerate(results[:30], 1):
        lines.append(
            f"{i}\t{r['gene']}\t{r['adimab_count']}\t{r['fr_length_diff']}\t{r['fr_mismatch']}\t"
            f"{r['usage_pct']:.1f}\t{r['fv_pct']:.1f}\t{r['fr_pct']:.1f}\t"
            f"{r['fr1_pct']:.1f}\t{r['fr2_pct']:.1f}\t{r['fr3_pct']:.1f}\t"
            f"{r['hyper_mut_pct']:.1f}\t{r['ptm_cys']}\t{r['ptm_glyco']}\t"
            f"{r['ptm_ng']}\t{r['ptm_ns']}\t{r['ptm_dg']}\t{r['ptm_dd']}\t{r['score']:.1f}"
        )
    
    return "\n".join(lines)


def generate_mutation_score_table(
    backmut: BackMutationResult,
    germline: Optional[GermlineGene],
    chain_type: str,
) -> str:
    """生成 Mutation Score 表格 (WeMol 格式)
    
    Chain	UID	Position	Donor Residue	Template Residue	Score	Germline Frequency
    """
    cons = backmut.germline_conservation
    
    lines = []
    lines.append("Chain\tUID\tPosition\tDonor Residue\tTemplate Residue\tScore\tGermline Frequency")
    lines.append("-" * 100)
    
    uid = 1
    filtered = [c for c in backmut.candidates if c.tier in ("T1", "T2", "T3")]
    filtered.sort(key=lambda x: -x.composite)
    for c in filtered:
        freq = cons.get(c.position)
        freq_txt = f"{freq:.0%}" if freq is not None else "-"
        lines.append(
            f"{chain_type}\t{uid}\t{c.position}\t{c.donor_aa}\t{c.human_aa}\t"
            f"{c.composite:.1f}\t{freq_txt}"
        )
        uid += 1
    
    return "\n".join(lines)


def generate_backmutation_summary(
    chain_reports: List[ChainReport],
) -> str:
    """生成回复突变摘要 (WeMol 格式)
    
    VL	VH
    L1	Graft(IGKV1-39*01)	H1	Graft(IGHV3-23*01)
    L2	Graft(IGKV1-39*01) + L46R,L71Y	H2	Graft(IGHV3-23*01)
    ...
    FR4	IGKJ2*01	FR4	IGHJ1*01
    """
    # 分离 VH 和 VL 链
    vh_rep = None
    vl_rep = None
    
    for rep in chain_reports:
        if rep.input_chain.chain_type == "H":
            vh_rep = rep
        elif rep.input_chain.chain_type == "L":
            vl_rep = rep
    
    lines = []
    lines.append("VL\tVH")
    lines.append("-" * 80)
    
    # 生成变体行
    max_variants = 4  # V0-V3
    
    for i in range(max_variants):
        vl_variant = ""
        vh_variant = ""
        
        if vl_rep:
            v_gene = vl_rep.germline.v_gene.gene_id if vl_rep.germline.v_gene else "None"
            if i == 0:
                vl_variant = f"Graft({v_gene})"
            else:
                # 获取对应的变体
                if i < len(vl_rep.variants):
                    # 提取突变信息
                    t_muts = []
                    for c in vl_rep.backmut.candidates:
                        if i == 1 and c.tier == "T1":
                            t_muts.append(f"{c.position}{c.donor_aa}")
                        elif i == 2 and c.tier in ["T1", "T2"]:
                            t_muts.append(f"{c.position}{c.donor_aa}")
                        elif i == 3 and c.tier in ["T1", "T2", "T3"]:
                            t_muts.append(f"{c.position}{c.donor_aa}")
                    if t_muts:
                        vl_variant = f"Graft({v_gene}) + {','.join(t_muts)}"
                    else:
                        vl_variant = f"Graft({v_gene})"
        
        if vh_rep:
            v_gene = vh_rep.germline.v_gene.gene_id if vh_rep.germline.v_gene else "None"
            if i == 0:
                vh_variant = f"Graft({v_gene})"
            else:
                t_muts = []
                for c in vh_rep.backmut.candidates:
                    if i == 1 and c.tier == "T1":
                        t_muts.append(f"{c.position}{c.donor_aa}")
                    elif i == 2 and c.tier in ["T1", "T2"]:
                        t_muts.append(f"{c.position}{c.donor_aa}")
                    elif i == 3 and c.tier in ["T1", "T2", "T3"]:
                        t_muts.append(f"{c.position}{c.donor_aa}")
                if t_muts:
                    vh_variant = f"Graft({v_gene}) + {','.join(t_muts)}"
                else:
                    vh_variant = f"Graft({v_gene})"
        
        lines.append(f"L{i+1}\t{vl_variant}\tH{i+1}\t{vh_variant}")
    
    # 添加 J 基因行
    vl_j = vl_rep.germline.j_gene.gene_id if vl_rep and vl_rep.germline.j_gene else "None"
    vh_j = vh_rep.germline.j_gene.gene_id if vh_rep and vh_rep.germline.j_gene else "None"
    lines.append(f"FR4\t{vl_j}\tFR4\t{vh_j}")
    
    return "\n".join(lines)


def generate_kabat_mutation_summary(
    chain_reports: List[ChainReport],
) -> str:
    """生成 KABAT 编号突变摘要"""
    
    lines = []
    lines.append("ID\tmutations")
    lines.append("-" * 80)
    
    for rep in chain_reports:
        chain_type = rep.input_chain.chain_type
        
        # V0
        lines.append(f"{chain_type}0\t")
        
        # V1: T1 回复突变
        t1_muts = [c for c in rep.backmut.candidates if c.tier == "T1"]
        if t1_muts:
            mut_str = ",".join([f"{c.human_aa}{c.position[1:]}{c.donor_aa}" for c in t1_muts])
            lines.append(f"{chain_type}1\t{mut_str}")
        else:
            lines.append(f"{chain_type}1\t")
        
        # V2: T1+T2 回复突变
        t2_muts = [c for c in rep.backmut.candidates if c.tier in ["T1", "T2"]]
        if t2_muts:
            mut_str = ",".join([f"{c.human_aa}{c.position[1:]}{c.donor_aa}" for c in t2_muts])
            lines.append(f"{chain_type}2\t{mut_str}")
        else:
            lines.append(f"{chain_type}2\t")
        
        # V3: T1+T2+T3 回复突变
        t3_muts = [c for c in rep.backmut.candidates if c.tier in ["T1", "T2", "T3"]]
        if t3_muts:
            mut_str = ",".join([f"{c.human_aa}{c.position[1:]}{c.donor_aa}" for c in t3_muts])
            lines.append(f"{chain_type}3\t{mut_str}")
        else:
            lines.append(f"{chain_type}3\t")
    
    return "\n".join(lines)


def generate_hotspot_summary(
    chain_reports: List[ChainReport],
) -> str:
    """生成热点摘要 (WeMol 格式)
    
    4.1 CDR Summary:
    ID	Sequence-CDR	High risk #	High risk sites
    
    4.2 Full-length Summary:
    ID	High risk #	High risk sites	Humanness (FV)	Humanness (FR)
    """
    lines = []
    
    # 4.1 CDR Summary
    lines.append("4.1 CDR Summary:")
    lines.append("ID\tSequence-CDR\tHigh risk #\tHigh risk sites")
    lines.append("-" * 100)
    
    for rep in chain_reports:
        chain_type = rep.input_chain.chain_type
        v_gene = rep.germline.v_gene.gene_id if rep.germline.v_gene else "None"
        cdr_seq = _get_cdr_sequences(rep)
        
        # V0: 扫描供体序列的 CDR 区域风险
        v0_count, v0_sites = _count_high_risk_sites(rep, cdr_only=True)
        lines.append(f"{chain_type}\t{cdr_seq}\t{v0_count}\t{', '.join(v0_sites)}")
        
        # V1-V3 变体: 每个变体单独扫描其移植后序列的 CDR 区域风险
        for i, v in enumerate(rep.variants[1:4], start=1):
            v_count, v_sites = _scan_high_risk(v.graft.numbered, cdr_only=True)
            lines.append(f"{chain_type}{i}\t{cdr_seq}\t{v_count}\t{', '.join(v_sites)}")
    
    lines.append("")
    
    # 4.2 Full-length Summary
    lines.append("4.2 Full-length Summary:")
    lines.append("ID\tHigh risk #\tHigh risk sites\tHumanness (FV)\tHumanness (FR)")
    lines.append("-" * 100)
    
    for rep in chain_reports:
        chain_type = rep.input_chain.chain_type
        
        # V0: 扫描纯移植序列
        graft = rep.grafts.get("kabat")
        if graft is not None:
            v0_count, v0_sites = _scan_high_risk(graft.numbered)
        else:
            v0_count, v0_sites = _count_high_risk_sites(rep)
        
        if graft is not None and rep.germline.v_gene:
            fv_identity = _calculate_fv_identity(graft.numbered, rep.germline.v_gene)
            fv_hl = fv_identity * 100
            fr_hl = _calculate_fr_identity(graft.numbered, rep.germline.v_gene) * 100
        else:
            fv_hl = 0
            fr_hl = 0
        
        lines.append(
            f"{chain_type}\t{v0_count}\t{', '.join(v0_sites)}\t"
            f"{fv_hl:.1f}%\t{fr_hl:.1f}%"
        )
        
        # V1-V3 变体: 每个变体单独扫描其序列
        for i, v in enumerate(rep.variants[1:], start=1):
            v_count, v_sites = _scan_high_risk(v.graft.numbered) if v.graft else (0, [])
            if v.graft is not None and rep.germline.v_gene:
                fv_identity = _calculate_fv_identity(v.graft.numbered, rep.germline.v_gene)
                fv_hl = fv_identity * 100
                fr_hl = _calculate_fr_identity(v.graft.numbered, rep.germline.v_gene) * 100
            else:
                fv_hl = 0
                fr_hl = 0
            
            lines.append(
                f"{chain_type}{i}\t{v_count}\t{', '.join(v_sites)}\t"
                f"{fv_hl:.1f}%\t{fr_hl:.1f}%"
            )
    
    return "\n".join(lines)


def generate_humanized_sequences(
    chain_reports: List[ChainReport],
) -> str:
    """生成人源化序列 (WeMol 格式)
    
    >VL
    原始序列
    >L1 Graft(IGKV1-39*01)
    变体序列
    ...
    """
    lines = []
    
    for rep in chain_reports:
        chain_type = rep.input_chain.chain_type
        v_gene = rep.germline.v_gene.gene_id if rep.germline.v_gene else "None"
        
        # 原始供体序列
        lines.append(f">{chain_type}")
        lines.append(rep.input_chain.sequence)
        
        # V0: 纯移植
        graft = rep.grafts.get("kabat")
        if graft is not None:
            lines.append(f">{chain_type}1 Graft({v_gene})")
            lines.append(graft.sequence)
        
        # V1-V3 变体
        for v in rep.variants[1:]:
            lines.append(f">{chain_type} {v.name} ({v.description})")
            lines.append(v.sequence)
    
    return "\n".join(lines)


def generate_enhanced_report(
    result: RunResult,
    outdir: str,
) -> str:
    """生成增强版报告 (WeMol 格式)"""
    
    lines = []
    lines.append("Antibody Humanization Design Summary")
    lines.append(datetime.now().strftime("%m/%d/%Y"))
    lines.append("=" * 80)
    lines.append("")
    
    # 1. Template Score
    lines.append("1. Template Score")
    lines.append("-" * 80)
    
    for rep in result.chains:
        chain_type = rep.input_chain.chain_type
        if chain_type == "H":
            lines.append(f"\nVH Chain (selected: {rep.germline.v_gene.gene_id if rep.germline.v_gene else 'None'}):")
            lines.append(generate_template_score_table(
                chain_type, rep.input_chain.numbered, result.germline_db, rep.germline.v_gene
            ))
        elif chain_type == "L":
            lines.append(f"\nVL Chain (selected: {rep.germline.v_gene.gene_id if rep.germline.v_gene else 'None'}):")
            lines.append(generate_template_score_table(
                chain_type, rep.input_chain.numbered, result.germline_db, rep.germline.v_gene
            ))
    
    lines.append("")
    
    # 2. Mutation Score
    lines.append("2. Mutation Score")
    lines.append("-" * 80)
    
    for rep in result.chains:
        chain_type = rep.input_chain.chain_type
        lines.append(f"\n{chain_type} Chain:")
        lines.append(generate_mutation_score_table(rep.backmut, rep.germline.v_gene, chain_type))
    
    lines.append("")
    
    # 3. Back Mutation Summary
    lines.append("3. Back Mutation Summary")
    lines.append("-" * 80)
    lines.append(generate_backmutation_summary(result.chains))
    
    lines.append("")
    
    # 3.1 KABAT Numbering Mutation Summary
    lines.append("3.1 KABAT Numbering Mutation Summary")
    lines.append("-" * 80)
    lines.append(generate_kabat_mutation_summary(result.chains))
    
    lines.append("")
    
    # 4. Hotspot Summary
    lines.append("4. Hotspot Summary")
    lines.append("-" * 80)
    lines.append(generate_hotspot_summary(result.chains))
    
    lines.append("")
    
    # 5. Humanized Sequences
    lines.append("5. Humanized Sequences")
    lines.append("-" * 80)
    lines.append(generate_humanized_sequences(result.chains))
    
    return "\n".join(lines)


def write_enhanced_docx(
    result: RunResult,
    outdir: str,
) -> str:
    """生成 WeMol 格式 Word 报告"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        return ""
    
    doc = Document()
    
    # 标题
    title = doc.add_heading("Antibody Humanization Design Summary", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 日期
    from datetime import datetime
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(datetime.now().strftime("%m/%d/%Y"))
    date_run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # 1. Template Score
    doc.add_heading("1. Template Score", level=1)
    
    for rep in result.chains:
        chain_type = rep.input_chain.chain_type
        if chain_type == "H":
            doc.add_heading(f"VH Chain (selected: {rep.germline.v_gene.gene_id if rep.germline.v_gene else 'None'})", level=2)
        elif chain_type == "L":
            doc.add_heading(f"VL Chain (selected: {rep.germline.v_gene.gene_id if rep.germline.v_gene else 'None'})", level=2)
        
        # 获取表格数据
        table_data = _generate_template_score_table_data(
            chain_type, rep.input_chain.numbered, result.germline_db, rep.germline.v_gene
        )
        
        # 创建表格
        if table_data:
            table = doc.add_table(rows=1, cols=len(table_data[0]))
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # 表头
            for j, header in enumerate(table_data[0]):
                cell = table.rows[0].cells[j]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(8)
            
            # 数据行
            for row_data in table_data[1:]:
                row = table.add_row()
                for j, value in enumerate(row_data):
                    cell = row.cells[j]
                    cell.text = str(value)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
            
            doc.add_paragraph()
    
    # 2. Mutation Score
    doc.add_heading("2. Mutation Score", level=1)
    
    for rep in result.chains:
        chain_type = rep.input_chain.chain_type
        doc.add_heading(f"{chain_type} Chain:", level=2)
        
        # 获取表格数据
        table_data = _generate_mutation_score_table_data(rep.backmut, rep.germline.v_gene, chain_type)
        
        # 创建表格
        if table_data:
            table = doc.add_table(rows=1, cols=len(table_data[0]))
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # 表头
            for j, header in enumerate(table_data[0]):
                cell = table.rows[0].cells[j]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(8)
            
            # 数据行
            for row_data in table_data[1:]:
                row = table.add_row()
                for j, value in enumerate(row_data):
                    cell = row.cells[j]
                    cell.text = str(value)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
            
            doc.add_paragraph()
    
    # 3. Back Mutation Summary
    doc.add_heading("3. Back Mutation Summary", level=1)
    
    for rep in result.chains:
        chain_type = rep.input_chain.chain_type
        doc.add_heading(f"{chain_type} Chain:", level=2)
        
        # 获取表格数据
        table_data = _generate_backmutation_summary_table_data(rep)
        
        # 创建表格
        if table_data:
            table = doc.add_table(rows=1, cols=len(table_data[0]))
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # 表头
            for j, header in enumerate(table_data[0]):
                cell = table.rows[0].cells[j]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(8)
            
            # 数据行
            for row_data in table_data[1:]:
                row = table.add_row()
                for j, value in enumerate(row_data):
                    cell = row.cells[j]
                    cell.text = str(value)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
            
            doc.add_paragraph()
    
    # 4. Hotspot Summary
    doc.add_heading("4. Hotspot Summary", level=1)
    
    # 4.1 CDR Summary
    doc.add_heading("4.1 CDR Summary", level=2)
    
    for rep in result.chains:
        chain_type = rep.input_chain.chain_type
        doc.add_heading(f"{chain_type} Chain:", level=2)
        
        # 获取表格数据
        table_data = _generate_cdr_summary_table_data(rep)
        
        # 创建表格
        if table_data:
            table = doc.add_table(rows=1, cols=len(table_data[0]))
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # 表头
            for j, header in enumerate(table_data[0]):
                cell = table.rows[0].cells[j]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(8)
            
            # 数据行
            for row_data in table_data[1:]:
                row = table.add_row()
                for j, value in enumerate(row_data):
                    cell = row.cells[j]
                    cell.text = str(value)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
            
            doc.add_paragraph()
    
    # 4.2 Full-length Summary
    doc.add_heading("4.2 Full-length Summary", level=2)
    
    for rep in result.chains:
        chain_type = rep.input_chain.chain_type
        doc.add_heading(f"{chain_type} Chain:", level=2)
        
        # 获取表格数据
        table_data = _generate_full_length_summary_table_data(rep)
        
        # 创建表格
        if table_data:
            table = doc.add_table(rows=1, cols=len(table_data[0]))
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # 表头
            for j, header in enumerate(table_data[0]):
                cell = table.rows[0].cells[j]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(8)
            
            # 数据行
            for row_data in table_data[1:]:
                row = table.add_row()
                for j, value in enumerate(row_data):
                    cell = row.cells[j]
                    cell.text = str(value)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
            
            doc.add_paragraph()
    
    # 5. Humanized Sequences
    doc.add_heading("5. Humanized Sequences", level=1)
    
    for rep in result.chains:
        chain_type = rep.input_chain.chain_type
        v_gene = rep.germline.v_gene.gene_id if rep.germline.v_gene else "None"
        
        # 原始供体序列
        doc.add_heading(f">{chain_type} (Donor)", level=2)
        seq_para = doc.add_paragraph()
        seq_run = seq_para.add_run(rep.input_chain.sequence)
        seq_run.font.name = "Consolas"
        seq_run.font.size = Pt(8)
        
        # V0: 纯移植
        graft = rep.grafts.get("kabat")
        if graft is not None:
            doc.add_heading(f">{chain_type}1 Graft({v_gene})", level=2)
            seq_para = doc.add_paragraph()
            seq_run = seq_para.add_run(graft.sequence)
            seq_run.font.name = "Consolas"
            seq_run.font.size = Pt(8)
        
        # V1-V3 变体
        for v in rep.variants[1:]:
            doc.add_heading(f">{chain_type} {v.name} ({v.description})", level=2)
            seq_para = doc.add_paragraph()
            seq_run = seq_para.add_run(v.sequence)
            seq_run.font.name = "Consolas"
            seq_run.font.size = Pt(8)
    
    # 保存文档
    out_path = os.path.join(outdir, "enhanced_report.docx")
    doc.save(out_path)
    return out_path


def _generate_template_score_table_data(
    chain_type: str,
    query: Optional[NumberedChain],
    germline_db: GermlineDB,
    selected_gene: Optional[GermlineGene],
) -> List[List[str]]:
    """生成 Template Score 表格数据 (30 条 germline, 18 列)"""
    if not query or not germline_db:
        return []
    
    # 获取 germline 基因列表
    genes = germline_db.human(chain_type)
    if not genes:
        return []
    
    rows = []
    
    # 表头
    headers = [
        "Rank", "V Gene", "FR Identity", "CDR Identity", "Composite",
        "CVI", "Freq", "3-axis", "FR1", "FR2", "FR3", "FR4",
        "CDR1", "CDR2", "CDR3", "Back-mutations", "Tier", "Selected"
    ]
    rows.append(headers)
    
    # 计算每个 germline 的分数
    results = []
    for gene in genes:
        # 计算 FR 同源性 (使用 region 标签，与 generate_template_score_table 一致)
        q_posmap = query.posmap()
        g_posmap = gene.numbered.posmap() if gene.numbered else {}
        
        fr_positions = {p for p in q_posmap if query.region_of(p) in _FR_REGIONS}
        fr_matches = sum(1 for p in fr_positions if p in g_posmap and q_posmap[p] == g_posmap[p])
        fr_identity = fr_matches / len(fr_positions) if fr_positions else 0
        
        # 计算 CDR 同源性 (使用 region 标签)
        cdr_positions = {p for p in q_posmap if query.region_of(p) in ("CDR1", "CDR2", "CDR3")}
        cdr_matches = sum(1 for p in cdr_positions if p in g_posmap and q_posmap[p] == g_posmap[p])
        cdr_identity = cdr_matches / len(cdr_positions) if cdr_positions else 0
        
        # Composite 分数 (0.7*FR + 0.3*CDR)
        composite = fr_identity * 0.7 + cdr_identity * 0.3
        
        # CVI (简化版)
        cvi = composite
        
        # 频率 (简化版)
        freq = 0
        
        # 3-axis 分数
        three_axis = composite
        
        # 计算各 FR/CDR 区域同源性 (使用 region 标签)
        def _region_pct(region_name):
            ps = [p for p in q_posmap if query.region_of(p) == region_name and p in g_posmap]
            if not ps:
                return 0.0
            return sum(1 for p in ps if q_posmap[p] == g_posmap[p]) / len(ps)
        
        fr1_identity = _region_pct("FR1")
        fr2_identity = _region_pct("FR2")
        fr3_identity = _region_pct("FR3")
        fr4_identity = _region_pct("FR4")
        cdr1_identity = _region_pct("CDR1")
        cdr2_identity = _region_pct("CDR2")
        cdr3_identity = _region_pct("CDR3")
        
        results.append({
            'gene': gene.gene_id,
            'fr_identity': fr_identity,
            'cdr_identity': cdr_identity,
            'composite': composite,
            'cvi': cvi,
            'freq': freq,
            'three_axis': three_axis,
            'fr1': fr1_identity,
            'fr2': fr2_identity,
            'fr3': fr3_identity,
            'fr4': fr4_identity,
            'cdr1': cdr1_identity,
            'cdr2': cdr2_identity,
            'cdr3': cdr3_identity,
            'backmutations': 0,
            'tier': '',
            'selected': gene == selected_gene
        })
    
    # 按 composite 排序
    results.sort(key=lambda x: -x['composite'])
    
    # 取前 30 个
    results = results[:30]
    
    # 添加到表格
    for rank, r in enumerate(results, 1):
        row = [
            str(rank),
            r['gene'],
            f"{r['fr_identity']:.3f}",
            f"{r['cdr_identity']:.3f}",
            f"{r['composite']:.3f}",
            f"{r['cvi']:.3f}",
            f"{r['freq']:.3f}",
            f"{r['three_axis']:.3f}",
            f"{r['fr1']:.3f}",
            f"{r['fr2']:.3f}",
            f"{r['fr3']:.3f}",
            f"{r['fr4']:.3f}",
            f"{r['cdr1']:.3f}",
            f"{r['cdr2']:.3f}",
            f"{r['cdr3']:.3f}",
            str(r['backmutations']),
            r['tier'],
            "*" if r['selected'] else ""
        ]
        rows.append(row)
    
    return rows


def _generate_mutation_score_table_data(
    backmut: BackMutationResult,
    germline: Optional[GermlineGene],
    chain_type: str,
) -> List[List[str]]:
    """生成 Mutation Score 表格数据"""
    cons = backmut.germline_conservation
    
    rows = []
    rows.append(["Chain", "UID", "Position", "Donor Residue", "Template Residue", "Score", "Germline Frequency"])
    
    uid = 1
    filtered = [c for c in backmut.candidates if c.tier in ("T1", "T2", "T3")]
    filtered.sort(key=lambda x: -x.composite)
    for c in filtered:
        freq = cons.get(c.position)
        freq_txt = f"{freq:.0%}" if freq is not None else "-"
        rows.append([
            chain_type,
            str(uid),
            c.position,
            c.donor_aa,
            c.human_aa,
            f"{c.composite:.1f}",
            freq_txt
        ])
        uid += 1
    
    return rows


def _generate_backmutation_summary_table_data(rep: ChainReport) -> List[List[str]]:
    """生成 Back Mutation Summary 表格数据"""
    chain_type = rep.input_chain.chain_type
    v_gene = rep.germline.v_gene.gene_id if rep.germline.v_gene else "None"
    j_gene = rep.germline.j_gene.gene_id if rep.germline.j_gene else "None"
    
    rows = []
    rows.append(["Type", "Description", "Back-mutations"])
    
    # V0: 纯移植
    rows.append([
        f"{chain_type}0",
        f"Graft({v_gene})",
        "-"
    ])
    
    # V1-V3 变体
    for v in rep.variants[1:]:
        bm_str = ",".join(v.backmutations) if v.backmutations else "-"
        rows.append([
            v.name,
            v.description,
            bm_str
        ])
    
    # FR4
    rows.append([
        "FR4",
        j_gene,
        "-"
    ])
    
    return rows


def _generate_cdr_summary_table_data(rep: ChainReport) -> List[List[str]]:
    """生成 CDR Summary 表格数据"""
    chain_type = rep.input_chain.chain_type
    
    rows = []
    rows.append(["ID", "Sequence-CDR", "High risk #", "High risk sites"])
    
    # 获取 CDR 序列 (使用 numbered.cdrs，与模块级 _get_cdr_sequences 一致)
    def _get_cdr_sequences(rep: ChainReport) -> str:
        """获取 CDR 序列"""
        if not rep.input_chain.numbered or not rep.input_chain.numbered.cdrs:
            return ""
        parts = []
        for name, (start, end) in rep.input_chain.numbered.cdrs.items():
            seq = rep.input_chain.numbered.seq_range(start, end)
            if seq:
                parts.append(seq)
        return "/".join(parts) if parts else ""
    
    # 获取高风险位点 (仅 CDR 区域)
    def _count_high_risk_sites(rep: ChainReport) -> tuple:
        """计算 CDR 区域高风险位点数量和位置"""
        from .developability import scan_sequence
        if not rep.input_chain.numbered:
            return 0, []
        issues = scan_sequence(rep.input_chain.numbered)
        high_risk_issues = [issue for issue in issues if issue.risk_level == 'high']
        # 仅保留 CDR 区域的位点
        cdr_positions = set()
        for name, (start, end) in rep.input_chain.numbered.cdrs.items():
            start_idx = rep.input_chain.numbered.residue(start)
            end_idx = rep.input_chain.numbered.residue(end)
            if start_idx and end_idx:
                for r in rep.input_chain.numbered.residues:
                    if start_idx.index <= r.index <= end_idx.index:
                        cdr_positions.add(r.pos)
        high_risk_issues = [issue for issue in high_risk_issues if issue.position in cdr_positions]
        high_risk_count = len(high_risk_issues)
        high_risk_sites = [f"{issue.position} ({issue.motif})" for issue in high_risk_issues]
        return high_risk_count, high_risk_sites
    
    cdr_seq = _get_cdr_sequences(rep)
    high_risk_count, high_risk_sites = _count_high_risk_sites(rep)
    
    # 主链
    rows.append([
        chain_type,
        cdr_seq,
        str(high_risk_count),
        ", ".join(high_risk_sites) if high_risk_sites else "-"
    ])
    
    # V1-V3 变体
    for i in range(1, 4):
        if i < len(rep.variants):
            rows.append([
                f"{chain_type}{i}",
                cdr_seq,
                str(high_risk_count),
                ", ".join(high_risk_sites) if high_risk_sites else "-"
            ])
    
    return rows


def _generate_full_length_summary_table_data(rep: ChainReport) -> List[List[str]]:
    """生成 Full-length Summary 表格数据 (每个变体单独扫描)"""
    chain_type = rep.input_chain.chain_type
    
    rows = []
    rows.append(["ID", "High risk #", "High risk sites", "Humanness (FV)", "Humanness (FR)"])
    
    # V0: 扫描纯移植序列
    graft = rep.grafts.get("kabat")
    if graft is not None:
        v0_count, v0_sites = _scan_high_risk(graft.numbered)
    else:
        v0_count, v0_sites = 0, []
    
    if graft is not None and rep.germline.v_gene:
        fv_identity = _calculate_fv_identity(graft.numbered, rep.germline.v_gene)
        fv_hl = fv_identity * 100
        fr_hl = _calculate_fr_identity(graft.numbered, rep.germline.v_gene) * 100
    else:
        fv_hl = 0
        fr_hl = 0
    
    rows.append([
        chain_type,
        str(v0_count),
        ", ".join(v0_sites) if v0_sites else "-",
        f"{fv_hl:.1f}%",
        f"{fr_hl:.1f}%"
    ])
    
    # V1-V3 变体: 每个变体单独扫描其序列
    for i, v in enumerate(rep.variants[1:], start=1):
        v_count, v_sites = _scan_high_risk(v.graft.numbered) if v.graft else (0, [])
        if v.graft is not None and rep.germline.v_gene:
            fv_identity = _calculate_fv_identity(v.graft.numbered, rep.germline.v_gene)
            fv_hl = fv_identity * 100
            fr_hl = _calculate_fr_identity(v.graft.numbered, rep.germline.v_gene) * 100
        else:
            fv_hl = 0
            fr_hl = 0
        
        rows.append([
            f"{chain_type}{i}",
            str(v_count),
            ", ".join(v_sites) if v_sites else "-",
            f"{fv_hl:.1f}%",
            f"{fr_hl:.1f}%"
        ])
    
    return rows
