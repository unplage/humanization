"""Professional Word (.docx) report generation.

Produces humanization_report.docx with:
  cover page, executive summary, input/format, numbering annotation,
  germline selection, CDR grafting (4 schemes), back-mutation tables,
  variant sequences (complete), minimal reversion & framework matrix,
  humanness metrics, developability scan, structural notes, experimental
  recommendations, and full-sequence appendices (parent / germline / graft /
  variants, plus a per-position Kabat map).

Requires python-docx (optional dependency; skipped gracefully when absent).
"""

from __future__ import annotations

import os
from datetime import date
from typing import Dict, List, Optional

from .config import TIER_LABELS
from .pipeline import ChainReport, RunResult

MONO = "Consolas"


def _mono(run, size=9):
    run.font.name = MONO
    run.font.size = docx_pt(size)
    return run


def docx_pt(v: int):
    from docx.shared import Pt
    return Pt(v)


def _num(v, nd=2):
    if v is None:
        return "-"
    return f"{v:.{nd}f}"


def _pos_key(p: str):
    return (p[0], int("".join(c for c in p if c.isdigit())), p)


def build_docx(result: RunResult, out_path: str) -> str:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()

    # ---- base styles ----
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    for hname, sz, col in [("Heading 1", 16, "1F4E79"), ("Heading 2", 13, "2E74B5"),
                           ("Heading 3", 11.5, "2E74B5")]:
        h = doc.styles[hname]
        h.font.name = "Calibri"
        h.font.size = Pt(sz)
        h.font.color.rgb = RGBColor.from_string(col)

    def para(text="", bold=False, italic=False, size=10.5, mono=False,
             align=None, space_after=6):
        p = doc.add_paragraph()
        if align:
            p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        r = p.add_run(text)
        r.bold, r.italic = bold, italic
        r.font.size = Pt(size)
        if mono:
            r.font.name = MONO
        return p

    def seq_para(label, seq, size=9):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{label}: ")
        r.bold = True
        r.font.size = Pt(size)
        r2 = p.add_run(seq)
        r2.font.name = MONO
        r2.font.size = Pt(size)
        return p

    def table(headers, rows, widths=None, mono_cols=()):
        t = doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Light Grid Accent 1"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, h in enumerate(headers):
            c = t.cell(0, j)
            c.text = ""
            r = c.paragraphs[0].add_run(str(h))
            r.bold = True
            r.font.size = Pt(9)
        for i, row in enumerate(rows):
            for j, v in enumerate(row):
                c = t.cell(i + 1, j)
                c.text = ""
                r = c.paragraphs[0].add_run(str(v))
                r.font.size = Pt(8.5)
                if j in mono_cols:
                    r.font.name = MONO
        if widths:
            for j, w in enumerate(widths):
                for row in t.rows:
                    row.cells[j].width = Inches(w)
        doc.add_paragraph()
        return t

    def h1(t): doc.add_heading(t, level=1)
    def h2(t): doc.add_heading(t, level=2)
    def h3(t): doc.add_heading(t, level=3)

    cfg = result.config

    # ============================ COVER ============================
    for _ in range(4):
        doc.add_paragraph()
    para("Antibody Humanization Report", bold=True, size=26,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    para("抗体人源化分析报告", bold=True, size=16,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    para(f"Report generated: {date.today().isoformat()}", size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    para(f"Format: {str(result.format).upper()} | "
         f"CDR scheme (variants): {cfg.cdr_scheme if cfg else 'kabat'}",
         size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    para("Antibody humanization pipeline - opencode/humanization",
         italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # ============================ EXECUTIVE SUMMARY ============================
    h1("1. Executive Summary (执行摘要)")
    for rep in result.chains:
        c = rep.input_chain
        v = rep.germline.v_gene
        j = rep.germline.j_gene
        h2(f"{c.name}  ({c.chain_type}{' | VHH' if c.is_vhh else ''})")
        if v:
            s = rep.germline.scores
            para(f"Chosen human germline: {v.gene_id} + {j.gene_id if j else '?'} "
                 f"(FR identity {_num(s.get('fr_identity'), 3)}, "
                 f"CDR identity {_num(s.get('cdr_identity'), 3)}, "
                 f"CVI {rep.cvi_homology:.3f})", size=10.5)
        tiers = {}
        for b in rep.backmut.candidates:
            tiers[b.tier] = tiers.get(b.tier, 0) + 1
        para(f"Back-mutation candidates: {sum(tiers.values())} total "
             f"({', '.join(f'{k}: {v}' for k, v in sorted(tiers.items()))})", size=10.5)
        hl = rep.human_likeness
        if hl:
            para("Human-likeness of pure graft: "
                 + ", ".join(f"{k} {v:.1f}%" for k, v in hl.items()), size=10.5)
        for vv in rep.variants:
            para(f"  {vv.name}: {vv.description} "
                 f"({len(vv.backmutations)} back-mutations)", size=10)
        rec = [vv for vv in rep.variants if vv.name.endswith("_V2")]
        if rec:
            para(f"Recommended production variant: {rec[0].name}", bold=True)
    if result.warnings:
        h3("Warnings")
        for w in result.warnings:
            para(f"- {w}", size=9.5, italic=True)
    doc.add_page_break()

    # ============================ 2. INPUT ============================
    h1("2. Input Sequences and Format Detection (输入与格式鉴定)")
    for rep in result.chains:
        c = rep.input_chain
        h2(f"{c.name} - {c.chain_type} chain"
           + (" (camelid VHH detected)" if c.is_vhh else ""))
        seq_para("Sequence", c.sequence)
        para(f"Length: {len(c.sequence)} aa | "
             f"VHH hallmark score: {c.vhh_score}/4", size=9.5)
        if c.warnings:
            for w in c.warnings:
                para(f"  warning: {w}", size=9, italic=True)

    # ============================ 3. NUMBERING ============================
    h1("3. Kabat Numbering and Region Annotation (编号与区域注释)")
    para("All positions in this report use Kabat numbering (anchor-based "
         "portable engine; cross-validated with ANARCI on the server).", size=9.5)
    for rep in result.chains:
        c = rep.input_chain
        h3(f"{c.name}")
        rows = []
        for region in ("FR1", "CDR1", "FR2", "CDR2", "FR3", "CDR3", "FR4"):
            rs = [r for r in c.numbered.residues if r.region == region]
            if rs:
                rows.append([region, f"{rs[0].pos}..{rs[-1].pos}", len(rs),
                             "".join(r.aa for r in rs)])
        table(["Region", "Positions", "Length", "Sequence"], rows, mono_cols=(3,))

    # ============================ 4. GERMLINE ============================
    h1("4. Human Germline Selection (人源 Germline 选择)")
    for rep in result.chains:
        c = rep.input_chain
        h3(f"{c.name}")
        if rep.germline.v_gene:
            v = rep.germline.v_gene
            j = rep.germline.j_gene
            para(f"V gene: {v.gene_id}", bold=True)
            seq_para("V gene sequence", v.sequence)
            if j:
                para(f"J gene: {j.gene_id}", bold=True)
                seq_para("J gene sequence", j.sequence)
            table(["Rank", "Gene", "FR identity", "CDR identity"],
                  [[i + 1, g.gene_id, _num(s.get("fr_identity"), 3),
                    _num(s.get("cdr_identity"), 3)]
                   for i, (g, s) in enumerate(rep.germline.alternatives[:5])],
                  widths=[0.6, 2.2, 1.2, 1.2])
        para(f"CVI homology (canonical+vernier+interface, BI 2024): "
             f"{rep.cvi_homology:.3f}", size=9.5)

    # ============================ 5. GRAFTING ============================
    h1("5. CDR Grafting (CDR 移植)")
    para("Four CDR boundary definitions are reported; the variant ladder "
         "uses the main scheme.", size=9.5)
    for rep in result.chains:
        c = rep.input_chain
        h3(f"{c.name}")
        for scheme, graft in rep.grafts.items():
            h4_txt = f"{scheme.upper()} scheme"
            p = doc.add_paragraph()
            r = p.add_run(h4_txt)
            r.bold = True
            r.font.size = Pt(10.5)
            cdr_txt = ", ".join(f"{k}: {v[0]}..{v[1]}"
                                for k, v in graft.numbered.cdrs.items())
            para(cdr_txt, size=9)
            seq_para("Grafted sequence", graft.sequence)
            donor = [p for p in graft.origin if graft.origin[p].startswith("donor")]
            para(f"Donor-derived positions ({len(donor)}): "
                 + ", ".join(sorted(donor, key=_pos_key)), size=8.5)

    # ============================ 6. BACK-MUTATIONS ============================
    h1("6. Back-mutation Analysis (回复突变分析)")
    for rep in result.chains:
        c = rep.input_chain
        h3(f"{c.name}")
        rows = []
        for b in sorted(rep.backmut.candidates, key=lambda x: -x.composite):
            emp = f"{b.empirical_ddG:+.2f}" if b.empirical_ddG is not None else "-"
            rows.append([b.position, b.donor_aa, b.human_aa, b.tier,
                         f"{b.composite:.1f}", emp,
                         "+".join(b.features) or "-"])
        table(["Pos", "Donor", "Human", "Tier", "Score", "ddG(emp)", "Features"],
              rows, widths=[0.7, 0.6, 0.6, 1.4, 0.7, 0.8, 2.4])
        # rationale detail
        h4 = doc.add_paragraph()
        r4 = h4.add_run("Per-position rationale (top candidates)")
        r4.bold = True
        r4.font.size = Pt(10.5)
        for b in sorted(rep.backmut.candidates, key=lambda x: -x.composite)[:20]:
            para(f"{b.position} {b.donor_aa}->{b.human_aa} [{b.tier}, "
                 f"{b.composite}]: " + "; ".join(b.rationale[:3]), size=8.5)

    # ============================ 7. VARIANTS ============================
    h1("7. Designed Variants (变体设计与完整序列)")
    for rep in result.chains:
        c = rep.input_chain
        h3(f"{c.name}")
        v0 = rep.variants[0]
        for v in rep.variants:
            para(f"{v.name} - {v.description}", bold=True, size=10.5)
            seq_para("Sequence", v.sequence, size=8.5)
            diffs = [p for p in v.graft.origin
                     if v.graft.origin[p] != v0.graft.origin.get(p)]
            if diffs:
                para("Changes vs V0: " + ", ".join(sorted(diffs, key=_pos_key)),
                     size=8.5)

    # ============================ 8. MINIMAL + MATRIX ============================
    h1("8. Minimal Reversion and Framework Matrix (最小回复集与框架矩阵)")
    for rep in result.chains:
        c = rep.input_chain
        h3(f"{c.name}")
        if rep.minimal_reversion:
            mr = rep.minimal_reversion
            para(f"Minimal reversion set ({mr.method}): "
                 + (", ".join(mr.positions) or "(none beyond graft)"), size=9.5)
            para(f"Contacts preserved: {mr.covered_contacts}/{mr.total_contacts} "
                 f"| {mr.note}", size=9, italic=True)
        if rep.matrix:
            table(["Alternative germline", "CVI", "# back-mutations (T1+T2)"],
                  [[e.germline.gene_id, f"{e.cvi:.3f}",
                    len(e.backmut.revert_positions(("T1", "T2")))]
                   for e in rep.matrix], widths=[2.6, 1.0, 2.0])

    # ============================ 9. HUMANNESS ============================
    h1("9. Humanness Assessment (人源化程度评估)")
    for rep in result.chains:
        c = rep.input_chain
        h3(f"{c.name}")
        hl = rep.human_likeness
        if hl:
            table(["CDR scheme", "FR germline identity (%)"],
                  [[k, f"{v:.1f}"] for k, v in hl.items()], widths=[2.0, 2.0])
        if rep.humanness:
            table(["Variant", "Sapiens mean", "OASis identity"],
                  [[k, h.get("sapiens_mean"), h.get("oasis_identity")]
                   for k, h in sorted(rep.humanness.items())],
                  widths=[2.6, 1.2, 1.2])

    # ============================ 10. DEVELOPABILITY ============================
    h1("10. Developability Risk Scan (可开发性风险扫描)")
    from .developability import scan_sequence
    for rep in result.chains:
        c = rep.input_chain
        h3(f"{c.name}")
        issues = scan_sequence(c.numbered)
        if issues:
            table(["Position", "Risk motif", "Context"],
                  [[i.position, i.motif, i.sequence_context] for i in issues],
                  widths=[1.0, 2.2, 2.4])
        else:
            para("No sequence-level developability risks detected.", size=9.5)
        # variants risks
        for v in rep.variants:
            vi = scan_sequence(v.graft.numbered)
            # Compare by (position, motif) — not motif alone — so a new
            # deamidation at position 95 is not masked by an existing one
            # at position 30.
            parent_pairs = {(i.position, i.motif) for i in issues}
            new_risks = [i for i in vi
                         if (i.position, i.motif) not in parent_pairs]
            if new_risks:
                para(f"{v.name} introduces: "
                     + ", ".join(f"{i.position} {i.motif}" for i in new_risks),
                     size=8.5, italic=True)

    # ============================ 11. STRUCTURE ============================
    if any(rep.structure_hints.data.get("buried") for rep in result.chains):
        h1("11. Structure-based Hints (AF3 结构提示)")
        for rep in result.chains:
            c = rep.input_chain
            h3(f"{c.name}")
            sh = rep.structure_hints.data
            buried = sh.get("buried") or {}
            cdr_contact = sh.get("cdr_contact") or {}
            ag_contact = sh.get("antigen_contact") or {}
            rows = []
            for pos in sorted(set(buried) | set(cdr_contact) | set(ag_contact),
                              key=_pos_key):
                rows.append([pos,
                             "yes" if buried.get(pos) else "-",
                             "yes" if cdr_contact.get(pos) else "-",
                             "yes" if ag_contact.get(pos) else "-"])
            table(["Position", "Buried", "CDR contact", "Antigen contact"],
                  rows, widths=[1.0, 0.9, 1.1, 1.4])

    # ============================ 12. EXPERIMENTAL ============================
    h1("12. Experimental Validation Recommendations (实验验证建议)")
    para("Follow docs/experimental_SOP.md. Key acceptance criteria:", size=10)
    table(["Criterion", "Go", "Conditional", "No-Go"],
          [["Affinity (KD vs parental)", "<= 2x", "2-5x (consider V1)", "> 5x"],
           ["Potency (EC50 vs parental)", "<= 3x", "3-5x", "> 5x"],
           ["SEC monomer", "> 95%", "90-95%", "< 90%"],
           ["Tm1 (DSF)", ">= parental - 3C", "-3..-6C", "< -6C"],
           ["Epitope binning", "same", "-", "changed"]],
          widths=[2.2, 1.0, 1.6, 1.0])
    para("After measuring affinities, run `humanize learn` with the KD data "
         "to recalibrate per-position scoring (docs/learning_loop.md).",
         size=9.5)

    # ============================ APPENDIX A: SEQUENCES ============================
    doc.add_page_break()
    h1("Appendix A. Complete Sequences (完整序列)")
    for rep in result.chains:
        c = rep.input_chain
        h2(f"{c.name}")
        seq_para("Parent (non-human)", c.sequence)
        if rep.germline.v_gene:
            seq_para(f"Human germline V ({rep.germline.v_gene.gene_id})",
                     rep.germline.v_gene.sequence)
        if rep.germline.j_gene:
            seq_para(f"Human germline J ({rep.germline.j_gene.gene_id})",
                     rep.germline.j_gene.sequence)
        for v in rep.variants:
            seq_para(v.name, v.sequence, size=8.5)

    # ============================ APPENDIX B: POSITION MAP ============================
    h1("Appendix B. Per-position Kabat Map (逐位点对照)")
    para("d = donor (non-human), g = human germline V, j = human J region.",
         size=9)
    for rep in result.chains:
        c = rep.input_chain
        h2(f"{c.name}")
        rows = []
        for r in c.numbered.residues:
            rows.append([r.pos, r.aa, r.region])
        table(["Position", "Parent aa", "Region"], rows, widths=[1.0, 1.0, 1.4],
              mono_cols=(0, 1))

    doc.save(out_path)
    return out_path


def write_docx_report(result: RunResult, outdir: str) -> Optional[str]:
    """Generate humanization_report.docx; returns path or None if
    python-docx is unavailable."""
    try:
        import docx  # noqa: F401
    except ImportError:
        return None
    path = os.path.join(outdir, "humanization_report.docx")
    return build_docx(result, path)
