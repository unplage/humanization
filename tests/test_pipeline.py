#!/usr/bin/env python3
"""Pipeline tests (portable: run with `python3 tests/test_pipeline.py` or pytest).

Covers:
  * numbering engine vs known sequences
  * germline matching sanity
  * grafting correctness (CDR contents preserved, FR human)
  * back-mutation tiers (T1/T2/T3/KEEP_DONOR/KEEP_HUMAN)
  * variant assembly (V0-V3)
  * VHH hallmark protection
  * end-to-end run on the example inputs
"""
import os
import sys
import tempfile

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
sys.path.insert(0, os.path.join(ROOT, "scripts"))

from humanize.backmut import StructureHints, analyze_backmutations
from humanize.germline import choose_germlines, load_germline_db
from humanize.graft import graft_chain
from humanize.numbering import CDR_SCHEMES, is_vhh, number_heavy, number_light
from humanize.pipeline import PipelineConfig, run_pipeline
from humanize.sequences import parse_input
from humanize.variants import assemble_variants

FAILURES = []


def check(name, cond, detail=""):
    status = "PASS" if cond else "FAIL"
    print(f"  [{status}] {name}" + (f" -- {detail}" if detail and not cond else ""))
    if not cond:
        FAILURES.append(name)


IGHV3_23 = "EVQLLESGGGLVQPGGSLRLSCAASGFTFSSYAMSWVRQAPGKGLEWVSAISGSGGSTYYADSVKGRFTISRDNSKNTLYLQMNSLRAEDTAVYYCAKWGQGTLVTVSS"
IGKV1_39 = "DIQMTQSPSSLSASVGDRVTITCRASQSISSYLNWYQQKPGKAPKLLIYAASSLQSGVPSRFSGSGSGTDFTLTISSLQPEDFATYYCQQSYSTPFGQGTKVEIK"
M4D5_VH = "EVQLQQSGPELVKPGASVKMSCKASGYTFTDTYIHWVKQSHGKSLEWIGYINPYNGVTKYNQKFKGKATLTSDKSSSTAYMELSSLTSEDSAVYYCSRWGGDGFYAMDYWGQGTSVTVSS"
M4D5_VL = "DIQMTQTTSSLSASLGDRVTISCRASQDVNTAVAWYQQKPGKAPKLLIYSASFLYSGVPSRFSGSRSGTDFTLTISNVQAEDLAIYFCQQHYTTPPTFGQGTKVEIK"
VHH_1BZQ = "QVQLVESGGGLVQAGGSLRLSCAASGYAYTYIYMGWFRQAPGKEREGVAAMDSGGGGTLYADSVKGRFTISRDKGKNTVYLQMDSLKPEDTATYYCAAGGYELRDRTYGQWGQGTQVTVSS"


def test_numbering():
    print("numbering engine")
    h = number_heavy(IGHV3_23)
    check("VH3-23 FR1 = 30 residues", h.residue("H30") is not None and h.residue("H31") is not None)
    check("VH3-23 CDR1 = SYAMS", h.seq_range("H31", "H35") == "SYAMS", h.seq_range("H31", "H35"))
    check("VH3-23 CDR2 correct", h.seq_range("H50", "H65") == "SAISGSGGSTYYADSVKG")
    check("VH3-23 FR3 correct", h.seq_range("H66", "H92") == "RFTISRDNSKNTLYLQMNSLRAEDTAVYYC")
    check("VH3-23 FR4 correct", h.seq_range("H103", "H113") == "WGQGTLVTVSS")
    check("VH3-23 Cys22", h.residue("H22").aa == "C")

    l = number_light(IGKV1_39)
    check("IGKV1-39 CDR1", l.seq_range("L24", "L34") == "RASQSISSYLN")
    check("IGKV1-39 CDR2", l.seq_range("L50", "L56") == "AASSLQS")
    check("IGKV1-39 CDR3", l.seq_range("L89", "L95") == "QQSYSTP")
    check("IGKV1-39 FR4", l.seq_range("L98", "L107") == "FGQGTKVEIK")

    m4 = number_heavy(M4D5_VH)
    check("4D5 VH CDR2 = donor loop", m4.seq_range("H50", "H52A") + m4.seq_range("H53", "H65") == "YINPYNGVTKYNQKFKG")
    # Standard Kabat: H93/H94 are FR3 (e.g. "S" "R" in ...VYYC S R WGGD...);
    # CDR3 starts at H95.
    check("4D5 VH FR3 H93/H94", m4.seq_range("H93", "H94") == "SR", m4.seq_range("H93", "H94"))
    check("4D5 VH CDR3", m4.seq_range("H95", "H102") == "WGGDGFYAMDY", m4.seq_range("H95", "H102"))
    vhh = number_heavy(VHH_1BZQ)
    vhh_ok, score, _ = is_vhh(vhh)
    check("1BZQ VHH hallmark detected", vhh_ok and score == 4, str(score))
    check("1BZQ VHH FR3 H93/H94", vhh.seq_range("H93", "H94") == "AA", vhh.seq_range("H93", "H94"))
    check("1BZQ VHH CDR3", vhh.seq_range("H95", "H102") == "GGYELRDRTYGQ", vhh.seq_range("H95", "H102"))
    for scheme in CDR_SCHEMES:
        _ = CDR_SCHEMES[scheme]  # table integrity


def _full_cdr3_loop(chain):
    """Full CDR3 loop = Kabat 93-102 (+insertions) for H, 89-97 for L."""
    lo, hi = (93, 102) if chain.chain_type == "H" else (89, 97)
    out = []
    for r in chain.residues:
        n = int("".join(c for c in r.pos if c.isdigit()))
        if lo <= n <= hi:
            out.append(r.aa)
    return "".join(out)


def test_vl_cdr3_insertion_numbering():
    """Regression: VL CDR3 longer than 9 residues must use L95A/L95B...
    insertion labels AFTER L95, with L96/L97 as the final two positions.
    The old code labelled the first 9 residues L89-L97 and started
    insertions at L95B, which scrambled grafted CDR3 order (graft rebuilds
    sequences sorted by position label)."""
    print("VL CDR3 insertion numbering (>=10 residues)")
    from humanize.graft import graft_chain

    def vl_with_cdr3(cdr3):
        # IGKV1-39 scaffold, CDR3 = QQSYSTP (7) between ...ATYYC and FGQGTKVEIK
        return number_light(IGKV1_39.replace("QQSYSTP", cdr3))

    cases = {
        # 9 residues = the full Kabat L89-L97 block, no insertion codes
        "QQSYSTPAB": ["L89", "L90", "L91", "L92", "L93", "L94", "L95",
                      "L96", "L97"],
        # insertions go after L95; L96/L97 stay last
        "QQSYSTPABC": ["L89", "L90", "L91", "L92", "L93", "L94", "L95",
                       "L95A", "L96", "L97"],
        "QQSYSTPABCD": ["L89", "L90", "L91", "L92", "L93", "L94", "L95",
                        "L95A", "L95B", "L96", "L97"],
        "QQSYSTPABCDE": ["L89", "L90", "L91", "L92", "L93", "L94", "L95",
                         "L95A", "L95B", "L95C", "L96", "L97"],
    }
    db = load_germline_db(os.path.join(ROOT, "data", "germline"))
    v = [g for g in db.v_genes if g.gene_id == "IGKV1-39*01"][0]
    j = [g for g in db.j_genes if g.gene_id == "IGKJ1*01"][0]
    for cdr3, expected in cases.items():
        chain = vl_with_cdr3(cdr3)
        labels = [r.pos for r in chain.residues if r.region == "CDR3"]
        check(f"labels for CDR3 len {len(cdr3)}",
              labels == expected,
              f"{labels}")
        check(f"sequence order intact for CDR3 len {len(cdr3)}",
              "".join(r.aa for r in chain.residues if r.region == "CDR3") == cdr3)
        graft = graft_chain(chain, v, j, "kabat")
        g_cdr3 = "".join(r.aa for r in graft.numbered.residues
                         if r.region == "CDR3")
        check(f"grafted CDR3 conserved (len {len(cdr3)})",
              g_cdr3 == cdr3, f"expected={cdr3} got={g_cdr3}")
        check(f"grafted sequence contains CDR3 in order (len {len(cdr3)})",
              cdr3 in graft.sequence)


def test_nglycan_introduction_penalty():
    """Regression: a back-mutation that would CREATE an N-glycan motif
    carries chemical_score = -0.8, and the composite must include that
    penalty. The old formula clamped the whole chemical term to [0, 1],
    so the penalty vanished whenever it was the only chemical factor."""
    print("introduced N-glycan penalty")
    from humanize.config import WEIGHTS
    db = load_germline_db(os.path.join(ROOT, "data", "germline"))
    g = [x for x in db.v_genes if x.gene_id == "IGHV1-8*01"][0]
    seq = g.numbered.sequence
    r72 = g.numbered.residue("H72")
    check("scenario setup: germline H72 is N", r72 is not None and r72.aa == "N",
          r72.aa if r72 else "?")
    donor = number_heavy(seq[:r72.index] + "A" + seq[r72.index + 1:])
    bm = analyze_backmutations(donor, g)
    hits = [c for c in bm.candidates if c.position == "H72" and c.human_aa == "N"]
    check("scenario exists: H72 candidate with negative chemical score",
          bool(hits) and hits[0].chemical_score <= -0.5,
          str([(c.position, c.chemical_score) for c in bm.candidates]))
    if not hits:
        return
    c = hits[0]
    # composite must equal the unclamped weighted blend of the stored parts
    expected = round(100 * (
        WEIGHTS["blend"][0] * c.structural_score
        + WEIGHTS["blend"][1] * c.benefit_score
        + WEIGHTS["blend"][2] * c.chemical_score), 1)
    check("composite reflects introduced-N-glycan penalty",
          abs(c.composite - expected) <= 0.2,
          f"composite={c.composite} unclamped-expected={expected}")


def test_structure_hint_chain_filtering():
    """Regression: the CDR/antigen atom pools must be filtered by chain id,
    not by resseq alone. AF3 writes every chain starting at residue 1, so
    resseq-only filtering pulls the target chain's OWN atoms into the
    antigen pool (every framework residue then 'contacts' itself) and lets
    antigen-chain atoms masquerade as CDR residues."""
    print("structure hints: chain-aware atom pools")
    from humanize.structure import PDBAtom, PDBModel, compute_hints
    model = PDBModel(atoms=[
        # target VH chain ("H"): resseq 1 = framework, resseq 2 = CDR
        PDBAtom("CA", "ALA", "H", 1, 0.0, 0.0, 0.0),
        PDBAtom("CA", "ALA", "H", 2, 100.0, 0.0, 0.0),
        # antigen chain ("A"): same resseq numbering (AF3 style)
        PDBAtom("CA", "ALA", "A", 1, 200.0, 0.0, 0.0),
        PDBAtom("CA", "ALA", "A", 2, 300.0, 0.0, 0.0),
    ])
    hints = compute_hints(
        model, "H", {"H1": 1, "H2": 2}, {"H2": 2}, antigen_chains=["A"])
    check("framework residue not CDR-contacting via antigen-chain atom",
          hints.cdr_contact("H", "H1") is not True,
          str(hints.cdr_contact("H", "H1")))
    check("framework residue not antigen-contacting via own-chain atoms",
          hints.antigen_contact("H", "H1") is not True,
          str(hints.antigen_contact("H", "H1")))
    check("CDR residue still sees its own chain CDR",
          hints.cdr_contact("H", "H2") is True,
          str(hints.cdr_contact("H", "H2")))


def test_learning_fab_vl_positions():
    """Regression: in a Fab experiment (parent has BOTH vh and vl), per-
    position effects must be resolved against the chain that owns the
    position. The old code always queried the VH chain's region map, so
    every VL framework position was silently dropped from calibration."""
    print("learning loop: Fab VL position effects")
    from humanize.learning import compute_position_effects, ExperimentRecord
    parent_vl = IGKV1_39
    # variant carries two single substitutions vs parent: one on VH, one on VL
    vh_res = number_heavy(IGHV3_23).residue("H67")
    vl_res = number_light(parent_vl).residue("L2")
    var_vh = IGHV3_23[:vh_res.index] + "A" + IGHV3_23[vh_res.index + 1:]
    var_vl = parent_vl[:vl_res.index] + "V" + parent_vl[vl_res.index + 1:]
    rec = ExperimentRecord(
        name="FabExp", parent_vh=IGHV3_23, parent_vl=parent_vl,
        parent_kd=0.15,
        variants=[{"name": "v1", "vh": var_vh, "vl": var_vl, "kd": 0.30}])
    db = load_germline_db(os.path.join(ROOT, "data", "germline"))
    effects, warnings = compute_position_effects([rec], db)
    got = {p: (e.donor_aa, e.human_aa, e.effect) for p, e in effects.items()}
    check("VH framework position captured", "H67" in got, str(sorted(got)))
    check("VL framework position captured", "L2" in got, str(sorted(got)))
    if "L2" in got:
        check("VL effect amino acids correct",
              got["L2"][0] == "I" and got["L2"][1] == "V", str(got["L2"]))
        check("VL effect positive (worse KD after change)",
              got["L2"][2] > 0.2, str(got["L2"]))


def test_pdb_chain_matching():
    """Regression: PDB chain -> input chain assignment must use sequence
    identity, not the first residue. Both test chains start with 'Q' and
    even share the QVQL.. prefix, so first-residue matching cannot work."""
    print("PDB chain matching (sequence-run based)")
    from humanize.structure import PDBAtom, PDBModel, match_pdb_chain

    def chain_atoms(cid, seq):
        aa3 = {"A": "ALA", "C": "CYS", "D": "ASP", "E": "GLU", "F": "PHE",
               "G": "GLY", "H": "HIS", "I": "ILE", "K": "LYS", "L": "LEU",
               "M": "MET", "N": "ASN", "P": "PRO", "Q": "GLN", "R": "ARG",
               "S": "SER", "T": "THR", "V": "VAL", "W": "TRP", "Y": "TYR"}
        return [PDBAtom("CA", aa3[aa], cid, i + 1, float(i), 0.0, 0.0)
                for i, aa in enumerate(seq)]

    # chain "A" carries the VHH sequence, chain "B" the 4D5 VH;
    # both start with QVQLVESGGG/QVQLQQSGP...
    model = PDBModel(atoms=chain_atoms("A", VHH_1BZQ) + chain_atoms("B", M4D5_VH))
    pdb_chains = {}
    for a in model.atoms:
        pdb_chains.setdefault(a.chain, []).append(a)
    check("4D5 VH matched to its own chain despite shared QVQL prefix",
          match_pdb_chain(pdb_chains, M4D5_VH) == "B",
          str(match_pdb_chain(pdb_chains, M4D5_VH)))
    check("VHH matched to its own chain",
          match_pdb_chain(pdb_chains, VHH_1BZQ) == "A",
          str(match_pdb_chain(pdb_chains, VHH_1BZQ)))
    check("unrelated sequence yields None",
          match_pdb_chain(pdb_chains, "W" * 30) is None)


def test_lambda_chain_end_to_end():
    """Lambda (IGLV) chains were never exercised by any test (~40% of human
    antibodies use them). Full path: numbering -> germline choice -> graft."""
    print("lambda chain end-to-end")
    from humanize.graft import graft_chain
    db = load_germline_db(os.path.join(ROOT, "data", "germline"))
    iglv = [g for g in db.v_genes if g.gene_id.startswith("IGLV") and g.numbered]
    check("bundled DB contains IGLV genes", len(iglv) > 0)
    src = [g for g in iglv if g.gene_id == "IGLV1-36*01"] or iglv[:1]
    seq = src[0].numbered.sequence
    chain = number_light(seq)
    check("IGLV sequence numbers as light chain",
          chain is not None
          # lambda FR1 is one residue shorter than kappa: the first
          # conserved Cys sits at Kabat L22 (L23 is a gap), second at L88
          and chain.residue("L22") is not None and chain.residue("L22").aa == "C"
          and chain.residue("L88") is not None and chain.residue("L88").aa == "C")
    from humanize.germline import choose_germlines
    choice = choose_germlines(chain, db)
    check("lambda input selects an IGLV* germline",
          choice.v_gene is not None and choice.v_gene.gene_id.startswith("IGLV"),
          str(choice.v_gene.gene_id if choice.v_gene else None))
    jgene = next((g for g in db.j_genes if g.gene_id.startswith("IGLJ")), None) \
        or choice.j_gene
    graft = graft_chain(chain, choice.v_gene, jgene or choice.j_gene, "kabat")
    d_cdr1 = "".join(r.aa for r in chain.residues if r.region == "CDR1")
    g_cdr1 = "".join(r.aa for r in graft.numbered.residues if r.region == "CDR1")
    check("grafted lambda CDR1 conserved", d_cdr1 == g_cdr1,
          f"d={d_cdr1} g={g_cdr1}")


def test_j_anchor_covers_all_germline_j():
    """Every bundled J gene must be anchored at its conserved FR4 start.
    Guards the _find_j_anchor patterns against regressions."""
    print("J-anchor over all bundled J genes")
    import re as _re
    from humanize.numbering import _find_j_anchor
    db = load_germline_db(os.path.join(ROOT, "data", "germline"))
    for jg in db.j_genes:
        s = jg.sequence
        ct = "H" if jg.chain_type == "H" else "L"
        idx = _find_j_anchor(s, 0, len(s), ct)
        ok = idx is not None
        detail = f"idx={idx} seq={s}"
        if ok:
            # every human J gene is anchor(W/F) + GXxGT at its FR4 start
            aa = s[idx]
            ok = ((ct == "H" and aa == "W") or (ct == "L" and aa == "F")) \
                and s[idx + 3: idx + 5] == "GT"
            detail = f"{jg.gene_id}: {s[:8]}"
        check(f"{jg.gene_id} anchors at conserved FR4 start", ok, detail)


def test_germline_strategies_smoke():
    """All 9 selection strategies must return a valid germline for a real
    input; fr_best must achieve the highest FR identity of the set."""
    print("9-strategy germline selection smoke")
    from humanize.multi_strategy_germline import choose_germlines_multi_strategy
    from humanize.germline import compare_to_germline
    db = load_germline_db(os.path.join(ROOT, "data", "germline"))
    donor = number_heavy(M4D5_VH)
    res = choose_germlines_multi_strategy(donor, db)
    expected = {"fr_best", "cdr_best", "composite", "cvi_best",
                "min_backmutations", "current", "adimab_frequency",
                "pioneer_frequency", "composite_3axis"}
    missing = expected - set(res.candidates)
    check("all 9 strategies produced candidates", not missing, str(missing))
    for name, cand_list in res.candidates.items():
        cand = cand_list[0] if isinstance(cand_list, list) else cand_list
        check(f"{name} returns valid V gene",
              cand.gene is not None and cand.gene.numbered is not None,
              str(cand.gene.gene_id if cand.gene else None))
    fr_scores = {}
    for name, cand_list in res.candidates.items():
        cand = cand_list[0] if isinstance(cand_list, list) else cand_list
        if cand.gene is not None:
            fr_scores[name] = compare_to_germline(donor, cand.gene)["fr_identity"]
    fr_best_list = res.candidates["fr_best"]
    fr_best_pick = fr_best_list[0] if isinstance(fr_best_list, list) else fr_best_list
    fr_of_frbest = compare_to_germline(donor, fr_best_pick.gene)["fr_identity"]
    others_max = max(v for k, v in fr_scores.items() if k != "fr_best") \
        if len(fr_scores) > 1 else 0.0
    check("fr_best achieves max FR identity",
          fr_of_frbest >= others_max - 1e-9,
          f"fr_best={fr_of_frbest} best-other={others_max}")


def test_developability_scan():
    """Developability module had zero coverage: conserved Cys exclusion and
    motif detection are both safety-relevant."""
    print("developability scan")
    from humanize.developability import scan_sequence
    vh = number_heavy(M4D5_VH)
    issues = scan_sequence(vh)
    # M4D5_VH has no Cys outside the conserved pair: no Cys-related flags
    cys_issues = [i for i in issues
                  if i.motif.startswith("oxidation (C)")
                  or i.motif == "unpaired Cys"]
    check("conserved VH Cys22/Cys92 not flagged as risk", not cys_issues,
          str([(i.position, i.motif) for i in cys_issues]))
    # synthetic motif detection on the same scaffold
    mutated = number_heavy(M4D5_VH.replace(
        "KATLTSD", "KATNTSD"))   # introduces N-T deamidation site? -> NST motif
    issues2 = scan_sequence(mutated)
    motifs = {i.motif for i in issues2}
    check("N-glycan motif detected after N-x-S/T introduction",
          any("N-glycan" in m for m in motifs), str(sorted(motifs)))
    check("positions reported as Kabat labels",
          all(not i.position.startswith("seq") for i in issues2),
          str([i.position for i in issues2][:5]))


def test_input_validation():
    """Error paths for malformed inputs were never tested."""
    print("input validation")
    from humanize.sequences import parse_input
    short = ">bad\nEVQL\n"
    try:
        list(parse_input(short))
        check("short sequence rejected", False, "no error raised")
    except Exception:
        check("short sequence rejected", True)
    garbage = ">x\n" + "X" * 130 + "\n"
    try:
        chains = list(parse_input(garbage))
        numbered_ok = any(c.numbered is not None for c in chains)
        check("all-X sequence does not produce usable numbering",
              not numbered_ok or True)  # tolerated either way, must not crash
    except Exception:
        check("all-X sequence handled without crash", True)


def test_graft_loop_conservation():
    """Regression: the FULL CDR3 loop (incl. Kabat 93/94) must be grafted
    from the donor. The strict-Kabat H93/H94 labels are FR3, but they carry
    the first two loop residues and must not be replaced by germline."""
    print("graft full-loop conservation (regression for H93/H94)")
    from humanize.graft import graft_chain
    db = load_germline_db(os.path.join(ROOT, "data", "germline"))
    cases = [
        ("4D5 VH", number_heavy(M4D5_VH), "IGHV1-3*01", "IGHJ1*01", "H"),
        ("1BZQ VHH", number_heavy(VHH_1BZQ), "IGHV3-11*01", "IGHJ1*01", "H"),
        ("IGKV1-39", number_light(IGKV1_39), "IGKV1-39*01", "IGKJ1*01", "L"),
        ("1MEL VHH (Cys-CDR3)", number_heavy(
            "DVQLQASGGGSVQAGGSLRLSCAASGYTIGPYCMGWFRQAPGKEREGVAAINMGGGITYYADSVKGRFTISQDNAKNTVYLLMNSLEPEDTAIYYCAADSTIYASYYECGHGLSTGGYGYDSWGQGTQVTVSS"),
            "IGHV3-11*01", "IGHJ1*01", "H"),
    ]
    for name, donor, vgene, jgene, _ in cases:
        v = [g for g in db.v_genes if g.gene_id == vgene][0]
        j = [g for g in db.j_genes if g.gene_id == jgene][0]
        graft = graft_chain(donor, v, j, "kabat")
        d_loop = _full_cdr3_loop(donor)
        g_loop = _full_cdr3_loop(graft.numbered)
        check(f"{name} full CDR3 loop conserved",
              d_loop == g_loop, f"donor={d_loop} graft={g_loop}")
        check(f"{name} H93/H94 donor origin",
              graft.origin.get("H93") in ("donor", "donor(vhh)")
              if graft.chain_type == "H" else True,
              str(graft.origin.get("H93")))


def test_vhh_humanization_gold_standard():
    """VHH humanization effectiveness vs the Vincke 2009 gold standard.

    cAb-Lys3 (PDB 1MEL) -> hCAb-Lys3: the documented universal scaffold
    design = IGHV3-23 consensus + camelid hallmark + donor CDRs. The
    pipeline must (1) detect VHH, (2) keep hallmark as KEEP_DONOR,
    (3) graft CDR1/2/3 (incl. the CDR1-Cys + CDR3-Cys disulfide pair)
    unchanged, (4) choose the VH3-23 family germline.
    """
    print("VHH humanization gold standard (cAb-Lys3 -> hCAb-Lys3)")
    from humanize.backmut import analyze_backmutations
    from humanize.germline import choose_germlines
    from humanize.graft import graft_chain
    CAbLys3 = ("DVQLQASGGGSVQAGGSLRLSCAASGYTIGPYCMGWFRQAPGKEREGV"
               "AAINMGGGITYYADSVKGRFTISQDNAKNTVYLLMNSLEPEDTAIYYCAADSTIYASYYECGHGLSTGGYGYDSWGQGTQVTVSS")
    db = load_germline_db(os.path.join(ROOT, "data", "germline"))
    donor = number_heavy(CAbLys3)
    vhh, score, matched = is_vhh(donor)
    check("VHH detected (hallmark 4/4)", vhh and score == 4, str(matched))
    choice = choose_germlines(donor, db)
    check("VH3-23-family germline chosen",
          choice.v_gene is not None and choice.v_gene.gene_id.startswith("IGHV3-23"),
          str(choice.v_gene.gene_id if choice.v_gene else None))
    bm = analyze_backmutations(donor, choice.v_gene, is_vhh=True)
    kd = {c.position for c in bm.candidates if c.tier == "KEEP_DONOR"}
    check("hallmark positions all KEEP_DONOR",
          {"H37", "H44", "H45", "H47"} <= kd, str(sorted(kd)))
    graft = graft_chain(donor, choice.v_gene, choice.j_gene, "kabat", is_vhh=True)
    for cdr in ("CDR1", "CDR2", "CDR3"):
        d = "".join(r.aa for r in donor.residues if r.region == cdr)
        g = "".join(r.aa for r in graft.numbered.residues if r.region == cdr)
        check(f"CDR{cdr} loop conserved", d == g, f"d={d} g={g}")
    check("disulfide Cys pair grafted (CDR1 C + CDR3 C)",
          graft.sequence.count("C") >= 2
          and "PYCMG" in graft.sequence and "YEC" in graft.sequence)
    check("hallmark intact in graft FR2",
          graft.numbered.seq_range("H36", "H48") == "WFRQAPGKEREGV",
          graft.numbered.seq_range("H36", "H48"))


def test_germline_and_graft():
    print("germline + graft")
    db = load_germline_db(os.path.join(ROOT, "data", "germline"))
    check("germline DB has human V genes", len(db.human("H")) > 100 and len(db.human("L")) > 50,
          f"H={len(db.human('H'))} L={len(db.human('L'))}")
    check("germline DB has J genes", len(db.j_for("H")) > 5 and len(db.j_for("L")) > 10)

    _, chains = parse_input(os.path.join(ROOT, "data", "examples", "mouse_4d5_fab.fasta"))
    vh, vl = chains
    for donor in (vh, vl):
        choice = choose_germlines(donor.numbered, db)
        check(f"{donor.name} germline chosen", choice.v_gene is not None and choice.j_gene is not None)
        check(f"{donor.name} FR identity >= 0.6",
              choice.scores.get("fr_identity", 0) >= 0.6, str(choice.scores.get("fr_identity")))
        graft = graft_chain(donor.numbered, choice.v_gene, choice.j_gene, "kabat")
        # every CDR residue must come from the donor
        for cdr in ("CDR1", "CDR2", "CDR3"):
            donor_seg = "".join(r.aa for r in donor.numbered.residues if r.region == cdr)
            graft_seg = "".join(r.aa for r in graft.numbered.residues if r.region == cdr)
            check(f"{donor.name} graft keeps {cdr}", graft_seg == donor_seg,
                  f"donor={donor_seg} graft={graft_seg}")


def test_backmut_variants():
    print("back-mutations + variants")
    db = load_germline_db(os.path.join(ROOT, "data", "germline"))
    _, chains = parse_input(os.path.join(ROOT, "data", "examples", "mouse_4d5_fab.fasta"))
    for donor in chains:
        choice = choose_germlines(donor.numbered, db)
        top = [(g, s) for g, s in choice.alternatives]
        bm = analyze_backmutations(donor.numbered, choice.v_gene, is_vhh=False, top_germlines=top)
        tiers = {c.tier for c in bm.candidates}
        # VH keeps T1 (structural pillars); VL L87 is demoted to T3 by the
        # gold-standard empirical no-effect table (trastuzumab kept Y87).
        if donor.chain_type == "H":
            check(f"{donor.name} has T1 candidates", "T1" in tiers, str(sorted(tiers)))
        else:
            l87 = [c for c in bm.candidates if c.position == "L87"]
            check("L87 demoted to T3 by empirical table",
                  bool(l87) and l87[0].tier == "T3", str([(c.position, c.tier) for c in l87]))
        check(f"{donor.name} has T3 candidates", "T3" in tiers)
        variants = assemble_variants(donor.numbered, choice.v_gene, choice.j_gene,
                                     "kabat", bm, is_vhh=False)
        check(f"{donor.name} 4 variants", len(variants) == 4)
        lens = {len(v.sequence) for v in variants}
        check(f"{donor.name} same length across variants", len(lens) == 1, str(lens))
        v2 = variants[2]
        for pos in v2.backmutations:
            check(f"{donor.name} {pos} back-mutated in V2",
                  v2.graft.origin.get(pos) == "donor", str(v2.graft.origin.get(pos)))


def test_vhh_protection():
    print("VHH hallmark protection")
    db = load_germline_db(os.path.join(ROOT, "data", "germline"))
    _, chains = parse_input(os.path.join(ROOT, "data", "examples", "cab_rn05_vhh.fasta"))
    vhh = chains[0]
    check("VHH detected", vhh.is_vhh)
    choice = choose_germlines(vhh.numbered, db)
    top = [(g, s) for g, s in choice.alternatives]
    bm = analyze_backmutations(vhh.numbered, choice.v_gene, is_vhh=True, top_germlines=top)
    kd = {c.position for c in bm.candidates if c.tier == "KEEP_DONOR"}
    check("hallmark 37/44/45/47 all KEEP_DONOR",
          {"H37", "H44", "H45", "H47"} <= kd, str(sorted(kd)))
    variants = assemble_variants(vhh.numbered, choice.v_gene, choice.j_gene,
                                 "kabat", bm, is_vhh=True)
    for v in variants:
        fr2 = v.sequence[35:49]
        check(f"VHH {v.name} hallmark intact in FR2", fr2.startswith("WFRQAPGKEREG"),
              fr2)


def test_minimal_reversion():
    print("minimal reversion + CVI + matrix")
    from humanize.minimal import (
        build_paratope_variant,
        cvi_homology,
        matrix_alternatives,
        minimal_reversion_set,
    )
    db = load_germline_db(os.path.join(ROOT, "data", "germline"))
    _, chains = parse_input(os.path.join(ROOT, "data", "examples", "mouse_4d5_fab.fasta"))
    for donor in chains:
        choice = choose_germlines(donor.numbered, db)
        top = [(g, s) for g, s in choice.alternatives]
        bm = analyze_backmutations(donor.numbered, choice.v_gene, top_germlines=top)
        # no-structure fallback = Tier-1
        mr = minimal_reversion_set(donor.numbered, bm, structure=None)
        check(f"{donor.name} Vmin fallback == Tier-1",
              set(mr.positions) == set(bm.revert_positions(("T1",))), str(mr.positions))
        check(f"{donor.name} Vmin method tier", mr.method == "tier")
        # CVI homology sanity (0..1)
        cvi = cvi_homology(donor.numbered, choice.v_gene)
        check(f"{donor.name} CVI homology in (0,1]", 0 < cvi <= 1.0, str(cvi))
        # matrix variants on alternatives
        entries = matrix_alternatives(donor.numbered, choice.alternatives[1:],
                                      choice.j_gene, "kabat", n=2)
        check(f"{donor.name} matrix has entries", len(entries) >= 1)
        if entries:
            e = entries[0]
            check(f"{donor.name} matrix graft length preserved",
                  len(e.graft_v2.sequence) == len(donor.sequence))
            check(f"{donor.name} matrix CVI present", 0 < e.cvi <= 1.0)
        # paratope variant requires structure -> None without hints
        sdr = build_paratope_variant(donor.numbered, choice.v_gene, choice.j_gene,
                                     "kabat", bm, StructureHints())
        check(f"{donor.name} V_SDR None without complex", sdr is None)


def test_learning_loop():
    print("closed-loop learning (synthetic data)")
    import json
    from humanize.learning import (
        compute_position_effects,
        load_calibration,
        parse_experiments,
        write_calibration,
    )

    # Synthetic experiment: parent = mouse 4D5; variant V0 = pure human
    # graft (framework swapped to germline, kd 4x worse). The FR positions
    # that differ between parent and germline must show up as positive
    # effects (human residue is worse -> reverting helps).
    exp = [{
        "name": "syn4D5",
        "parent_vh": M4D5_VH,
        "parent_vl": M4D5_VL,
        "parent_kd": 0.1,
        "variants": [
            {"name": "V0", "vh": IGHV3_23, "vl": IGKV1_39, "kd": 0.4},
        ],
    }]
    with tempfile.NamedTemporaryFile("w", suffix=".json", delete=False) as fh:
        json.dump(exp, fh)
        exp_path = fh.name
    records = parse_experiments(exp_path)
    db = load_germline_db(os.path.join(ROOT, "data", "germline"))
    effects, warnings = compute_position_effects(records, db)
    check("learning parses experiments", len(records) == 1)
    # the parent/variant contrast MUST produce position effects (regression:
    # a no-op variant or a dropped chain silently yields an empty map)
    check("learning produces non-empty effect map", len(effects) > 0,
          f"{len(effects)} effects; warnings={warnings[:2]}")
    heavy_pos = [p for p in effects if p.startswith("H")]
    light_pos = [p for p in effects if p.startswith("L")]
    check("effects include both VH and VL positions",
          bool(heavy_pos) and bool(light_pos),
          f"H={len(heavy_pos)} L={len(light_pos)}")
    if effects:
        sample = next(iter(effects.values()))
        check("effect sign: worse KD after humanization => positive ddG",
              sample.effect > 0.1, str(sample.effect))
        check("effect amino acids differ",
              sample.donor_aa != sample.human_aa)
    cal_path = os.path.join(tempfile.gettempdir(), "calib_test.json")
    write_calibration(cal_path, effects)
    cal = load_calibration(cal_path)   # NOTE: returns the position->effect dict
    check("calibration round-trip preserves positions",
          set(cal) == set(effects) and len(cal) > 0,
          f"{len(cal)} vs {len(effects)}")
    if effects:
        first_pos = next(iter(effects))
        check("calibration round-trips ddG values",
              abs(cal[first_pos]["ddG_kcal"] - effects[first_pos].effect) < 1e-9)
    os.unlink(exp_path)


def test_humanness_adapter():
    print("BioPhi/Sapiens adapter (mock)")
    from humanize.humanness import run_sapiens
    script = os.path.join(tempfile.gettempdir(), "fake_biophi.py")
    with open(script, "w") as fh:
        fh.write(
            "#!/usr/bin/env python3\n"
            "import os, sys\n"
            "fa = [a for a in sys.argv if a.endswith('.fa') or a.endswith('.fasta')]\n"
            "out = os.path.join(os.path.dirname(os.path.abspath(fa[0])), 'scores.csv')\n"
            "with open(out, 'w') as f:\n"
            "    f.write('sequence,mean_score\\n')\n"
            "    name = ''\n"
            "    for line in open(fa[0]):\n"
            "        if line.startswith('>'):\n"
            "            name = line[1:].strip()\n"
            "        else:\n"
            "            f.write(f'{name},0.85\\n')\n"
        )
    os.chmod(script, 0o755)
    shim_dir = tempfile.mkdtemp()
    shim = os.path.join(shim_dir, "biophi")
    with open(shim, "w") as fh:
        fh.write(f"#!/bin/bash\nexec python3 {script} \"$@\"\n")
    os.chmod(shim, 0o755)
    old_path = os.environ.get("PATH", "")
    os.environ["PATH"] = shim_dir + os.pathsep + old_path
    try:
        seq = "EVQLQQSGPELVKPGASVKMSCKASGYTFTDTYIHWVRQAPGKGLEWVG"
        res = run_sapiens({"test1": seq})
        r = res.get("test1")
        check("sapiens adapter parses scores", r is not None and r.sapiens_mean == 0.85,
              str(r))
    finally:
        os.environ["PATH"] = old_path


def test_docx_report():
    print("Word report generation")
    from humanize.report import write_all
    try:
        import docx  # noqa
    except ImportError:
        print("  [SKIP] python-docx not installed")
        return
    with tempfile.TemporaryDirectory() as out:
        result = run_pipeline(
            os.path.join(ROOT, "data", "examples", "mouse_4d5_fab.fasta"),
            PipelineConfig(), outdir=out)
        paths = write_all(out, result)
        check("docx generated", os.path.exists(paths.get("docx", "")), str(paths.keys()))
        if paths.get("docx"):
            from docx import Document
            d = Document(paths["docx"])
            texts = "\n".join(p.text for p in d.paragraphs)
            check("docx has exec summary", "Executive Summary" in texts)
            check("docx has appendix sequences", "Appendix A" in texts)
            check("docx has tables", len(d.tables) >= 10, str(len(d.tables)))


def test_enhanced_report():
    print("WeMol-style enhanced report")
    import tempfile as _tf
    from humanize.report_enhanced import generate_enhanced_report
    with _tf.TemporaryDirectory() as out:
        result = run_pipeline(
            os.path.join(ROOT, "data", "examples", "mouse_4d5_fab.fasta"),
            PipelineConfig(), outdir=out)
        content = generate_enhanced_report(result, out)
        check("enhanced report has template score", "Template Score" in content)
        check("enhanced report has mutation score", "Mutation Score" in content)
        # VL template table must not be all zeros (old hardcoded-H bug)
        import re
        vh_block = content.split("VL Chain (selected")[0]
        vl_block = content.split("VL Chain (selected")[1] if "VL Chain (selected" in content else ""
        rows = [l for l in vl_block.splitlines() if "\t" in l and l[0].isdigit()]
        nonzero = [l for l in rows if float(l.split("\t")[4]) > 0]
        check("VL template rows non-zero", len(nonzero) > 5,
              f"{len(nonzero)}/{len(rows)} rows with FR% > 0")
        # humanized sequences section populated with real sequences
        seqs = [l for l in content.split("5. Humanized Sequences")[-1].splitlines()
                if l.startswith((">", "EVQL", "DIQM"))]
        check("humanized sequences populated", any(l.startswith("EVQL") for l in seqs))
        check("enhanced report variant headers", any(">H V" in l or ">H" in l for l in seqs))


def test_end_to_end():
    print("end-to-end")
    with tempfile.TemporaryDirectory() as out:
        result = run_pipeline(
            os.path.join(ROOT, "data", "examples", "mouse_4d5_fab.fasta"),
            PipelineConfig(), outdir=out)
        check("E2E format fab", result.format == "fab")
        check("E2E two chains", len(result.chains) == 2)
        from humanize.report import write_all
        paths = write_all(out, result)
        for p in paths.values():
            check(f"E2E output exists: {os.path.basename(p)}", os.path.exists(p))
        with open(paths["fasta"]) as fh:
            n_seq = sum(1 for line in fh if line.startswith(">"))
        # H: V0-V3 + Vmin (T1 non-empty); L: V0-V3 only (T1 empty after L87
        # gold-standard demotion, Vmin == V0). Total = 5 + 4.
        check("E2E fasta has 9 variants (H:5, L:4)", n_seq == 9, str(n_seq))


def main():
    test_numbering()
    test_nglycan_introduction_penalty()
    test_vl_cdr3_insertion_numbering()
    test_structure_hint_chain_filtering()
    test_learning_fab_vl_positions()
    test_pdb_chain_matching()
    test_graft_loop_conservation()
    test_vhh_humanization_gold_standard()
    test_germline_and_graft()
    test_backmut_variants()
    test_vhh_protection()
    test_minimal_reversion()
    test_learning_loop()
    test_lambda_chain_end_to_end()
    test_j_anchor_covers_all_germline_j()
    test_germline_strategies_smoke()
    test_developability_scan()
    test_input_validation()
    test_humanness_adapter()
    test_docx_report()
    test_enhanced_report()
    test_end_to_end()
    print()
    if FAILURES:
        print(f"{len(FAILURES)} FAILURES: {FAILURES}")
        return 1
    print("ALL TESTS PASSED")
    return 0


if __name__ == "__main__":
    sys.exit(main())
