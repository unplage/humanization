#!/usr/bin/env python3
"""Generate USAN Standard Guide Word Document."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def create_document():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('抗体人源化评估标准指南', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('基于 USAN/WHO 标准和 IMGT 编号系统')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.add_paragraph()
    
    # ============================================================
    # Section 1: USAN Standard Overview
    # ============================================================
    doc.add_heading('1. USAN 标准简介', level=1)
    
    doc.add_heading('1.1 什么是 USAN', level=2)
    doc.add_paragraph(
        'USAN (United States Adopted Name) 是美国采用名称系统，由美国医学协会 (AMA) 和美国药典 (USP) '
        '联合管理。对于治疗性抗体，USAN 制定了命名规则，其中包含了人源化程度的评估标准。'
    )
    
    doc.add_heading('1.2 USAN 抗体命名规则', level=2)
    
    # Table 1: Naming rules
    table1 = doc.add_table(rows=4, cols=4)
    table1.style = 'Table Grid'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['命名后缀', '人源化程度', 'FR 身份要求', '示例']
    for i, header in enumerate(headers):
        cell = table1.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    data = [
        ['-zumab', '人源化 (Humanized)', 'FR ≥ 85%', 'Trastuzumab (曲妥珠单抗)'],
        ['-xi-', '嵌合 (Chimeric)', 'FR 70-85%', 'Rituximab (利妥昔单抗)'],
        ['-o-', '鼠源 (Murine)', 'FR < 70%', 'Muromonab (莫罗单抗)'],
    ]
    
    for row_idx, row_data in enumerate(data, 1):
        for col_idx, cell_text in enumerate(row_data):
            table1.rows[row_idx].cells[col_idx].text = cell_text
    
    doc.add_paragraph()
    
    doc.add_heading('1.3 USAN 是否是金标准', level=2)
    doc.add_paragraph('是的，USAN 是抗体人源化评估的金标准之一，但需要注意：')
    
    # Sub-items
    p = doc.add_paragraph()
    p.add_run('官方权威性').bold = True
    doc.add_paragraph('USAN 是 FDA 批准的治疗性抗体命名的官方标准', style='List Bullet')
    doc.add_paragraph('所有在美国上市的抗体药物必须遵循 USAN 命名规则', style='List Bullet')
    doc.add_paragraph('是监管机构（FDA、EMA、NMPA）认可的评估方法', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('评估基础').bold = True
    doc.add_paragraph('基于 IMGT 编号系统（国际免疫遗传学信息系统）', style='List Bullet')
    doc.add_paragraph('评估 FR 身份（框架区同源性），不包括 CDR', style='List Bullet')
    doc.add_paragraph('使用 WHO/INN 国际标准', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('局限性').bold = True
    doc.add_paragraph('仅评估 FR 身份：不考虑 CDR 身份', style='List Bullet')
    doc.add_paragraph('静态阈值：85% 的阈值是经验值，不是绝对标准', style='List Bullet')
    doc.add_paragraph('不考虑结构：仅基于序列比对，不考虑三维结构', style='List Bullet')
    doc.add_paragraph('不考虑功能：不评估结合亲力、稳定性等', style='List Bullet')
    
    doc.add_paragraph()
    
    # ============================================================
    # Section 2: Evaluation Methods
    # ============================================================
    doc.add_heading('2. 评估方法', level=1)
    
    doc.add_heading('2.1 重链和轻链分别评估', level=2)
    doc.add_paragraph('是的，必须分别评估：')
    
    # Table 2: Evaluation targets
    table2 = doc.add_table(rows=4, cols=3)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers2 = ['评估对象', '评估内容', '说明']
    for i, header in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    data2 = [
        ['重链 (VH)', 'FR 同源性、CDR 同源性、CVI', '单独评估，单独报告'],
        ['轻链 (VL)', 'FR 同源性、CDR 同源性、CVI', '单独评估，单独报告'],
        ['整体 Fv', '综合评估', '可选，但学术论文通常分别报告'],
    ]
    
    for row_idx, row_data in enumerate(data2, 1):
        for col_idx, cell_text in enumerate(row_data):
            table2.rows[row_idx].cells[col_idx].text = cell_text
    
    doc.add_paragraph()
    doc.add_paragraph('为什么分别评估：')
    doc.add_paragraph('重链和轻链可能匹配不同的 germline 基因', style='List Bullet')
    doc.add_paragraph('重链和轻链的人源化程度可能不同', style='List Bullet')
    doc.add_paragraph('监管机构要求分别提供数据', style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('2.2 CDR 分区规则', level=2)
    doc.add_paragraph('推荐使用 IMGT 定义：')
    
    # Table 3: CDR definitions
    table3 = doc.add_table(rows=7, cols=4)
    table3.style = 'Table Grid'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers3 = ['CDR', 'IMGT 位置', 'Kabat 位置', '学术论文推荐']
    for i, header in enumerate(headers3):
        cell = table3.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    data3 = [
        ['CDR-H1', '26-35 (11 残基)', '31-35 (5 残基)', 'IMGT ✓'],
        ['CDR-H2', '56-65 (10 残基)', '50-65 (16 残基)', 'IMGT ✓'],
        ['CDR-H3', '105-117 (13 残基)', '95-102 (8 残基)', 'IMGT ✓'],
        ['CDR-L1', '24-34 (11 残基)', '24-34 (11 残基)', 'IMGT ✓'],
        ['CDR-L2', '50-56 (7 残基)', '50-56 (7 残基)', 'IMGT ✓'],
        ['CDR-L3', '89-97 (9 残基)', '89-97 (9 残基)', 'IMGT ✓'],
    ]
    
    for row_idx, row_data in enumerate(data3, 1):
        for col_idx, cell_text in enumerate(row_data):
            table3.rows[row_idx].cells[col_idx].text = cell_text
    
    doc.add_paragraph()
    doc.add_paragraph('学术论文中的表述：')
    p = doc.add_paragraph()
    p.style = 'Quote'
    p.add_run('"CDR 定义采用 IMGT 编号系统 (Lefranc et al., 1999)。"')
    
    doc.add_paragraph()
    
    doc.add_heading('2.3 评估 FR 区还是整个 Fv 区', level=2)
    doc.add_paragraph('USAN 标准：评估 FR1+FR2+FR3（排除 FR4）')
    
    # Table 4: Evaluation regions
    table4 = doc.add_table(rows=5, cols=3)
    table4.style = 'Table Grid'
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers4 = ['评估区域', '包含内容', '学术论文推荐']
    for i, header in enumerate(headers4):
        cell = table4.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    data4 = [
        ['FR 区 (USAN)', 'FR1 + FR2 + FR3（排除 FR4）', 'USAN 标准要求 ✓'],
        ['FR4 区域', 'FR4（来自人 J 基因）', '单独报告，不用于命名'],
        ['Fv 区', 'FR + CDR', '可选，但不用于命名'],
        ['CDR 区', 'CDR1 + CDR2 + CDR3', '单独报告，不用于命名'],
    ]
    
    for row_idx, row_data in enumerate(data4, 1):
        for col_idx, cell_text in enumerate(row_data):
            table4.rows[row_idx].cells[col_idx].text = cell_text
    
    doc.add_paragraph()
    doc.add_paragraph('为什么排除 FR4：')
    doc.add_paragraph('FR4 来自人 J 基因，不是人 V 基因', style='List Bullet')
    doc.add_paragraph('人有 6 个 J 基因（IGHJ1-6），FR4 序列差异较大', style='List Bullet')
    doc.add_paragraph('监管机构（FDA/EMA）在审批时更关注 FR1+FR2+FR3 的同源性', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('USAN 标准的评估逻辑：')
    p = doc.add_paragraph()
    p.style = 'Quote'
    p.add_run('FR 同源性 = (FR1 相同残基 + FR2 相同残基 + FR3 相同残基) / (FR1 总残基 + FR2 总残基 + FR3 总残基)')
    
    doc.add_paragraph()
    
    # ============================================================
    # Section 3: Academic Best Practices
    # ============================================================
    doc.add_heading('3. 学术宣传最佳实践', level=1)
    
    doc.add_heading('3.1 论文/海报格式', level=2)
    
    p = doc.add_paragraph()
    p.add_run('摘要部分：').bold = True
    p = doc.add_paragraph()
    p.style = 'Quote'
    p.add_run('"该抗体采用 CDR 移植技术进行人源化，基于 IMGT 编号系统评估，重链 FR 同源性为 85.2%，轻链 FR 同源性为 92.3%，符合 USAN/WHO 人源化抗体标准。"')
    
    p = doc.add_paragraph()
    p.add_run('方法部分：').bold = True
    p = doc.add_paragraph()
    p.style = 'Quote'
    p.add_run('"人源化程度评估采用 USAN/WHO 标准 (WHO, 2012)，使用 IMGT 编号系统 (Lefranc et al., 1999)。FR 同源性定义为框架区 (FR1+FR2+FR3，排除 FR4) 中与人 germline 相同残基的比例。"')
    
    p = doc.add_paragraph()
    p.add_run('结果部分：').bold = True
    p = doc.add_paragraph()
    p.style = 'Quote'
    p.add_run('"表 X 展示了人源化评估结果。根据 USAN 标准 (FR ≥ 85% = 人源化)，该抗体被归类为人源化抗体 (-zumab)。"')
    
    doc.add_paragraph()
    
    doc.add_heading('3.2 结果报告示例', level=2)
    
    # Table 5: Example results
    table5 = doc.add_table(rows=3, cols=5)
    table5.style = 'Table Grid'
    table5.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers5 = ['链', 'Germline', 'FR 同源性 (IMGT)', 'CDR 同源性 (IMGT)', '人源化程度']
    for i, header in enumerate(headers5):
        cell = table5.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    data5 = [
        ['VH', 'IGHV1-46*01', '85.2%', '62.5%', 'Humanized (-zumab)'],
        ['VL', 'IGKV1-39*01', '92.3%', '78.9%', 'Humanized (-zumab)'],
    ]
    
    for row_idx, row_data in enumerate(data5, 1):
        for col_idx, cell_text in enumerate(row_data):
            table5.rows[row_idx].cells[col_idx].text = cell_text
    
    doc.add_paragraph()
    
    doc.add_heading('3.3 补充报告（可选）', level=2)
    
    # Table 6: Supplementary reports
    table6 = doc.add_table(rows=4, cols=3)
    table6.style = 'Table Grid'
    table6.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers6 = ['评估维度', '方法', '说明']
    for i, header in enumerate(headers6):
        cell = table6.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    data6 = [
        ['CVI 同源性', 'BI 2024 标准', 'Canonical + Vernier + Interface 区域'],
        ['Germline 频率', 'Wemol 方法', '人 germline 使用频率'],
        ['结构验证', 'AF3/晶体结构', '埋藏/CDR 接触分析'],
    ]
    
    for row_idx, row_data in enumerate(data6, 1):
        for col_idx, cell_text in enumerate(row_data):
            table6.rows[row_idx].cells[col_idx].text = cell_text
    
    doc.add_paragraph()
    
    # ============================================================
    # Section 4: Summary
    # ============================================================
    doc.add_heading('4. 总结', level=1)
    
    # Table 7: Summary
    table7 = doc.add_table(rows=5, cols=2)
    table7.style = 'Table Grid'
    table7.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers7 = ['问题', '推荐答案']
    for i, header in enumerate(headers7):
        cell = table7.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    data7 = [
        ['评估标准', 'USAN/WHO 标准 + IMGT 编号'],
        ['重链/轻链', '分别评估，分别报告'],
        ['CDR 分区', '使用 IMGT 定义'],
        ['评估区域', 'FR1+FR2+FR3（排除 FR4）'],
    ]
    
    for row_idx, row_data in enumerate(data7, 1):
        for col_idx, cell_text in enumerate(row_data):
            table7.rows[row_idx].cells[col_idx].text = cell_text
    
    doc.add_paragraph()
    
    # ============================================================
    # Section 5: Fully Human Antibodies
    # ============================================================
    doc.add_heading('5. 全人源抗体（Fully Human Antibody）', level=1)
    
    doc.add_heading('5.1 定义', level=2)
    doc.add_paragraph(
        '全人源抗体是指其可变区（VH 和 VL）序列 100% 来自人类抗体库，'
        '不含有任何鼠源或非人源序列的抗体。'
    )
    
    # Table 8: Antibody types comparison
    table8 = doc.add_table(rows=5, cols=4)
    table8.style = 'Table Grid'
    table8.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers8 = ['抗体类型', '来源', '鼠源序列比例', '命名后缀']
    for i, header in enumerate(headers8):
        cell = table8.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    data8 = [
        ['鼠源抗体', '小鼠杂交瘤', '100%', '-omab'],
        ['嵌合抗体', '鼠 V 区 + 人 C 区', '~30%', '-ximab'],
        ['人源化抗体', '鼠 CDR + 人 FR', '~5-10%', '-zumab'],
        ['全人源抗体', '人 V 区 + 人 C 区', '0%', '-umab'],
    ]
    
    for row_idx, row_data in enumerate(data8, 1):
        for col_idx, cell_text in enumerate(row_data):
            table8.rows[row_idx].cells[col_idx].text = cell_text
    
    doc.add_paragraph()
    
    doc.add_heading('5.2 命名规则', level=2)
    doc.add_paragraph('USAN/WHO 标准：全人源抗体使用 -umab 后缀')
    
    # Table 9: Naming examples
    table9 = doc.add_table(rows=5, cols=3)
    table9.style = 'Table Grid'
    table9.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers9 = ['命名后缀', '抗体类型', '示例']
    for i, header in enumerate(headers9):
        cell = table9.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    data9 = [
        ['-umab', '全人源 (Fully Human)', 'Adalimumab (阿达木单抗)'],
        ['-zumab', '人源化 (Humanized)', 'Trastuzumab (曲妥珠单抗)'],
        ['-ximab', '嵌合 (Chimeric)', 'Rituximab (利妥昔单抗)'],
        ['-omab', '鼠源 (Murine)', 'Muromonab (莫罗单抗)'],
    ]
    
    for row_idx, row_data in enumerate(data9, 1):
        for col_idx, cell_text in enumerate(row_data):
            table9.rows[row_idx].cells[col_idx].text = cell_text
    
    doc.add_paragraph()
    
    doc.add_heading('5.3 生成技术', level=2)
    
    p = doc.add_paragraph()
    p.add_run('转基因小鼠技术：').bold = True
    doc.add_paragraph('敲除小鼠的鼠源抗体基因，插入人源抗体基因片段', style='List Bullet')
    doc.add_paragraph('用抗原免疫小鼠，筛选人源抗体', style='List Bullet')
    doc.add_paragraph('代表技术：XenoMouse, HuMAb-Mouse, VelocImmune', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('噬菌体展示技术：').bold = True
    doc.add_paragraph('构建人源抗体文库（10^9-10^11 个克隆）', style='List Bullet')
    doc.add_paragraph('用抗原筛选结合噬菌体', style='List Bullet')
    doc.add_paragraph('代表技术：HuCAL, n-CoDeR, OrthoLibrary', style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('5.4 已批准的全人源抗体', level=2)
    
    # Table 10: Approved fully human antibodies
    table10 = doc.add_table(rows=8, cols=5)
    table10.style = 'Table Grid'
    table10.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers10 = ['药物名称', '靶点', '适应症', '批准年份', '技术']
    for i, header in enumerate(headers10):
        cell = table10.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    data10 = [
        ['Adalimumab', 'TNF-α', '类风湿关节炎', '2002', '转基因小鼠'],
        ['Natalizumab', 'α4 整合素', '多发性硬化', '2004', '噬菌体展示'],
        ['Denosumab', 'RANKL', '骨质疏松', '2010', '转基因小鼠'],
        ['Nivolumab', 'PD-1', '黑色素瘤', '2014', '转基因小鼠'],
        ['Pembrolizumab', 'PD-1', '黑色素瘤', '2014', '噬菌体展示'],
        ['Atezolizumab', 'PD-L1', '膀胱癌', '2016', '转基因小鼠'],
        ['Avelumab', 'PD-L1', 'Merkel 细胞癌', '2017', '噬菌体展示'],
    ]
    
    for row_idx, row_data in enumerate(data10, 1):
        for col_idx, cell_text in enumerate(row_data):
            table10.rows[row_idx].cells[col_idx].text = cell_text
    
    doc.add_paragraph()
    
    doc.add_heading('5.5 与人源化抗体的区别', level=2)
    
    # Table 11: Comparison
    table11 = doc.add_table(rows=6, cols=3)
    table11.style = 'Table Grid'
    table11.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers11 = ['方面', '全人源抗体', '人源化抗体']
    for i, header in enumerate(headers11):
        cell = table11.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    data11 = [
        ['V 区来源', '人抗体库', '鼠 CDR + 人 FR'],
        ['FR 同源性', '≥ 95%', '85-95%'],
        ['CDR 同源性', '≥ 90%', '50-80%'],
        ['免疫原性', '最低', '低'],
        ['命名后缀', '-umab', '-zumab'],
    ]
    
    for row_idx, row_data in enumerate(data11, 1):
        for col_idx, cell_text in enumerate(row_data):
            table11.rows[row_idx].cells[col_idx].text = cell_text
    
    doc.add_paragraph()
    
    # ============================================================
    # Section 6: IMGT vs Kabat
    # ============================================================
    doc.add_heading('6. IMGT vs Kabat：评估与工程化的双编号系统', level=1)
    
    doc.add_heading('6.1 历史背景', level=2)
    
    # Table 12: History
    table12 = doc.add_table(rows=3, cols=4)
    table12.style = 'Table Grid'
    table12.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers12 = ['系统', '创建时间', '创建者', '初始目的']
    for i, header in enumerate(headers12):
        cell = table12.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    data12 = [
        ['Kabat', '1979', 'Elvin Kabat', '序列比对和结构研究'],
        ['IMGT', '1989', 'Marie-Paule Lefranc', '免疫遗传学数据库标准化'],
    ]
    
    for row_idx, row_data in enumerate(data12, 1):
        for col_idx, cell_text in enumerate(row_data):
            table12.rows[row_idx].cells[col_idx].text = cell_text
    
    doc.add_paragraph()
    
    doc.add_heading('6.2 核心差异', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Kabat 编号系统：').bold = True
    doc.add_paragraph('CDR 边界：基于序列可变性（高变区）', style='List Bullet')
    doc.add_paragraph('位置数量：CDR 长度可变（5-26 残基）', style='List Bullet')
    doc.add_paragraph('插入编码：使用字母（H52A, H100B）', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('IMGT 编号系统：').bold = True
    doc.add_paragraph('CDR 边界：基于结构（β-折叠）', style='List Bullet')
    doc.add_paragraph('位置数量：CDR 长度固定（11-10-13 残基）', style='List Bullet')
    doc.add_paragraph('插入编码：使用数字（H52, H52A, H52B）', style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('6.3 为什么评估用 IMGT？', level=2)
    
    # Table 13: IMGT advantages
    table13 = doc.add_table(rows=5, cols=2)
    table13.style = 'Table Grid'
    table13.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers13 = ['优势', '说明']
    for i, header in enumerate(headers13):
        cell = table13.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    data13 = [
        ['标准化', '所有抗体使用统一的编号'],
        ['可比性', '不同抗体可以直接比较'],
        ['数据库', 'IMGT 数据库使用 IMGT 编号'],
        ['监管认可', 'USAN/WHO 标准采用 IMGT'],
    ]
    
    for row_idx, row_data in enumerate(data13, 1):
        for col_idx, cell_text in enumerate(row_data):
            table13.rows[row_idx].cells[col_idx].text = cell_text
    
    doc.add_paragraph()
    
    doc.add_heading('6.4 为什么工程化用 Kabat？', level=2)
    
    # Table 14: Kabat advantages
    table14 = doc.add_table(rows=5, cols=2)
    table14.style = 'Table Grid'
    table14.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers14 = ['优势', '说明']
    for i, header in enumerate(headers14):
        cell = table14.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    data14 = [
        ['结构相关性', 'Kabat CDR 定义与抗体三维结构更相关'],
        ['功能重要性', '高变区直接对应功能位点'],
        ['工程经验', '40+ 年的工程实践积累'],
        ['灵活性', '字母插入编码更适合 CDR 长度变化'],
    ]
    
    for row_idx, row_data in enumerate(data14, 1):
        for col_idx, cell_text in enumerate(row_data):
            table14.rows[row_idx].cells[col_idx].text = cell_text
    
    doc.add_paragraph()
    
    doc.add_heading('6.5 实际工作流程', level=2)
    
    doc.add_paragraph('1. 人源化评估 (IMGT)：使用 IMGT 编号评估 FR/CDR 同源性', style='List Number')
    doc.add_paragraph('2. CDR 移植 (Kabat)：使用 Kabat 编号识别 CDR 位置', style='List Number')
    doc.add_paragraph('3. 回复突变 (Kabat)：使用 Kabat 编号设计回复突变', style='List Number')
    doc.add_paragraph('4. 亲和力成熟 (Kabat)：使用 Kabat 编号定位 CDR 突变位点', style='List Number')
    
    doc.add_paragraph()
    
    doc.add_heading('6.6 总结', level=2)
    
    # Table 15: Summary comparison
    table15 = doc.add_table(rows=6, cols=3)
    table15.style = 'Table Grid'
    table15.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers15 = ['方面', 'IMGT', 'Kabat']
    for i, header in enumerate(headers15):
        cell = table15.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    data15 = [
        ['用途', '评估、比较、数据库', '工程化、设计、实验'],
        ['CDR 定义', '基于结构（β-折叠）', '基于序列可变性'],
        ['标准化', '高（统一编号）', '低（实验室差异）'],
        ['工程经验', '少（30+ 年）', '多（40+ 年）'],
        ['文献支持', '新文献', '大量历史文献'],
    ]
    
    for row_idx, row_data in enumerate(data15, 1):
        for col_idx, cell_text in enumerate(row_data):
            table15.rows[row_idx].cells[col_idx].text = cell_text
    
    doc.add_paragraph()
    doc.add_paragraph('结论：评估用 IMGT（标准化、监管认可），工程化用 Kabat（结构相关性、实验经验），两者互补。')
    
    doc.add_paragraph()
    
    # ============================================================
    # Section 7: References
    # ============================================================
    doc.add_heading('7. 参考文献', level=1)
    
    doc.add_paragraph('WHO. (2012). International Nonproprietary Names (INN) for Biological Products. World Health Organization.', style='List Number')
    doc.add_paragraph('Lefranc, M. P., et al. (1999). IMGT, the international ImMunoGeneTics database. Nucleic Acids Research, 27(1), 209-212.', style='List Number')
    doc.add_paragraph('Chothia, C., et al. (1989). Conformations of immunoglobulin hypervariable regions. Nature, 342(6245), 877-883.', style='List Number')
    doc.add_paragraph('Kabat, E. A., et al. (1991). Sequences of Proteins of Immunological Interest. 5th ed. US Department of Health and Human Services.', style='List Number')
    
    doc.add_paragraph()
    
    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('生成时间：2026年9月')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    return doc


if __name__ == '__main__':
    doc = create_document()
    output_path = '/home/jiemiaoxing/work/opencode_task/humanization/USAN标准指南.docx'
    doc.save(output_path)
    print(f"✅ Word document generated: {output_path}")
